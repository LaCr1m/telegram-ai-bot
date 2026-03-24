import os
import re
import logging
import httpx
import base64
import asyncio
import io
import json
import sqlite3
from datetime import date, timedelta, datetime
from zoneinfo import ZoneInfo
from tavily import TavilyClient
import openpyxl
from docx import Document as DocxDocument
try:
    from pypdf import PdfReader
except ImportError:
    from PyPDF2 import PdfReader
from telegram import Update, InlineKeyboardButton, InlineKeyboardMarkup
from telegram.ext import (
    ApplicationBuilder, MessageHandler, CommandHandler,
    CallbackQueryHandler, filters, ContextTypes,
)

# ── Logging ───────────────────────────────────────────────────────────────────

logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s [%(levelname)s] %(name)s: %(message)s",
    datefmt="%Y-%m-%d %H:%M:%S",
)
log = logging.getLogger("jarvis")

# ── Env vars ──────────────────────────────────────────────────────────────────

def _require_env(key: str) -> str:
    val = os.environ.get(key)
    if not val:
        log.warning("Env var %s is not set", key)
    return val or ""

TELEGRAM_TOKEN     = _require_env("TELEGRAM_TOKEN")
OPENROUTER_API_KEY = _require_env("OPENROUTER_API_KEY")
GROQ_API_KEY       = _require_env("GROQ_API_KEY")
TAVILY_API_KEY     = _require_env("TAVILY_API_KEY")
CF_API_TOKEN       = _require_env("CF_API_TOKEN")
GOOGLE_API_KEY = _require_env("GOOGLE_API_KEY")   # chat_id куди слати щоденний бриф

if not TELEGRAM_TOKEN:
    raise RuntimeError("TELEGRAM_TOKEN is required")

# ── URLs & models ─────────────────────────────────────────────────────────────

OPENROUTER_URL            = "https://openrouter.ai/api/v1/chat/completions"
GROQ_URL                  = "https://api.groq.com/openai/v1/chat/completions"
GROQ_WHISPER_URL          = "https://api.groq.com/openai/v1/audio/transcriptions"
GEMINI_URL         = "https://generativelanguage.googleapis.com/v1beta/models"
GEMINI_PRO_MODEL          = "gemini-2.0-flash"
GEMINI_FLASH_MODEL        = "gemini-1.5-flash-latest"
OPENROUTER_MODEL          = "meta-llama/llama-3.3-70b-instruct:free"
OPENROUTER_MODEL_FALLBACK = "meta-llama/llama-3.1-8b-instruct:free"
GROQ_MODEL                = "llama-3.3-70b-versatile"
VISION_MODEL              = "openrouter/auto"
CF_IMAGE_URL              = "https://api.cloudflare.com/client/v4/accounts/{account_id}/ai/run/@cf/black-forest-labs/flux-1-schnell"

# ── Constants ─────────────────────────────────────────────────────────────────

MAX_HISTORY_MESSAGES  = 20
SUMMARY_THRESHOLD     = 16
STYLE_UPDATE_INTERVAL = 10
MSG_CHUNK_SIZE        = 4000
MAX_VIDEO_SIZE        = 20 * 1024 * 1024
MAX_ARTICLE_CHARS     = 6000
MAX_DOC_PREVIEW_CHARS = 4000
CTX_DESCRIPTION_LEN   = 500
SEARCH_RESULTS        = 7
FLUX_STEPS            = 4
BRIEF_HOUR            = 9   # година щоденного брифу (Київ)

DB_PATH         = os.environ.get("DB_PATH", "/data/jarvis.db")
BLOCKED_DOMAINS = {"olx.ua", "olx.com.ua"}
_HEADERS        = {"User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36"}

TZ = ZoneInfo("Europe/Kyiv")

def now_kyiv() -> datetime:
    return datetime.now(TZ)

# ── SQLite storage ────────────────────────────────────────────────────────────

def _db() -> sqlite3.Connection:
    os.makedirs(os.path.dirname(DB_PATH), exist_ok=True)
    con = sqlite3.connect(DB_PATH, check_same_thread=False)
    con.row_factory = sqlite3.Row
    return con

def init_db() -> None:
    with _db() as con:
        con.executescript("""
            CREATE TABLE IF NOT EXISTS memory (
                user_id INTEGER PRIMARY KEY,
                data    TEXT NOT NULL DEFAULT '{}'
            );
            CREATE TABLE IF NOT EXISTS history (
                user_id  INTEGER PRIMARY KEY,
                messages TEXT NOT NULL DEFAULT '[]'
            );
            CREATE TABLE IF NOT EXISTS tasks (
                user_id INTEGER PRIMARY KEY,
                data    TEXT NOT NULL DEFAULT '[]'
            );
            CREATE TABLE IF NOT EXISTS reminders (
                id      INTEGER PRIMARY KEY AUTOINCREMENT,
                chat_id INTEGER NOT NULL,
                text    TEXT NOT NULL,
                fire_at TEXT NOT NULL
            );
            CREATE TABLE IF NOT EXISTS session_log (
                id         INTEGER PRIMARY KEY AUTOINCREMENT,
                user_id    INTEGER NOT NULL,
                created_at TEXT NOT NULL,
                summary    TEXT NOT NULL,
                mood       TEXT
            );
        """)
    log.info("DB initialized at %s", DB_PATH)

def save_session_log(user_id: int, summary: str, mood: str = "") -> None:
    with _db() as con:
        con.execute(
            "INSERT INTO session_log(user_id, created_at, summary, mood) VALUES(?,?,?,?)",
            (user_id, now_kyiv().isoformat(), summary[:500], mood),
        )
    # Зберігаємо лише останні 30 записів на користувача
    with _db() as con:
        con.execute(
            "DELETE FROM session_log WHERE user_id=? AND id NOT IN "
            "(SELECT id FROM session_log WHERE user_id=? ORDER BY id DESC LIMIT 30)",
            (user_id, user_id),
        )

def get_session_log(user_id: int, limit: int = 5) -> list[dict]:
    with _db() as con:
        rows = con.execute(
            "SELECT created_at, summary, mood FROM session_log "
            "WHERE user_id=? ORDER BY id DESC LIMIT ?",
            (user_id, limit),
        ).fetchall()
    return [dict(r) for r in rows]

# memory
def get_user_memory(user_id: int) -> dict:
    with _db() as con:
        row = con.execute("SELECT data FROM memory WHERE user_id=?", (user_id,)).fetchone()
    return json.loads(row["data"]) if row else {}

def update_user_memory(user_id: int, data: dict) -> None:
    mem = get_user_memory(user_id)
    mem.update(data)
    with _db() as con:
        con.execute(
            "INSERT INTO memory(user_id,data) VALUES(?,?) ON CONFLICT(user_id) DO UPDATE SET data=excluded.data",
            (user_id, json.dumps(mem, ensure_ascii=False)),
        )

# history
def load_history_db(user_id: int) -> list:
    with _db() as con:
        row = con.execute("SELECT messages FROM history WHERE user_id=?", (user_id,)).fetchone()
    return json.loads(row["messages"]) if row else []

def save_history_db(user_id: int, messages: list) -> None:
    with _db() as con:
        con.execute(
            "INSERT INTO history(user_id,messages) VALUES(?,?) ON CONFLICT(user_id) DO UPDATE SET messages=excluded.messages",
            (user_id, json.dumps(messages, ensure_ascii=False)),
        )

# tasks
def get_user_tasks(user_id: int) -> list:
    with _db() as con:
        row = con.execute("SELECT data FROM tasks WHERE user_id=?", (user_id,)).fetchone()
    return json.loads(row["data"]) if row else []

def set_user_tasks(user_id: int, tasks: list) -> None:
    with _db() as con:
        con.execute(
            "INSERT INTO tasks(user_id,data) VALUES(?,?) ON CONFLICT(user_id) DO UPDATE SET data=excluded.data",
            (user_id, json.dumps(tasks, ensure_ascii=False)),
        )

# reminders
def load_reminders() -> list:
    with _db() as con:
        rows = con.execute("SELECT id, chat_id, text, fire_at FROM reminders").fetchall()
    return [dict(r) for r in rows]

def add_reminder(chat_id: int, text: str, fire_at: datetime) -> int:
    with _db() as con:
        cur = con.execute(
            "INSERT INTO reminders(chat_id,text,fire_at) VALUES(?,?,?)",
            (chat_id, text, fire_at.isoformat()),
        )
    return cur.lastrowid

def delete_reminder(rid: int) -> None:
    with _db() as con:
        con.execute("DELETE FROM reminders WHERE id=?", (rid,))

# ── Emotion tones ─────────────────────────────────────────────────────────────

EMOTION_TONES = {
    "sad":     "Людина сумує. Визнай її почуття одним теплим реченням, не поспішай з порадами.",
    "angry":   "Людина роздратована. Не сперечайся. Дай відчути що її почули, потім переходь до фактів.",
    "anxious": "Людина тривожиться. Спочатку заспокій одним реченням. Структуруй відповідь чітко.",
    "happy":   "Людина радісна. Підтримай її енергію, будь живим.",
    "neutral": "",
}

# ── Prompts & personalities ───────────────────────────────────────────────────

_BASE_LANGUAGE_RULES = (
    "Завжди відповідай виключно українською мовою, незалежно від мови запиту. "
    "Використовуй грамотну українську без суржику і русизмів. "
    "ЗАБОРОНЕНО символи інших мов. Ніколи не вигадуй факти. "
    "ЗАБОРОНЕНО повторювати питання користувача — одразу відповідай по суті."
)

_CHARACTER_CORE = (
    "Ти J.A.R.V.I.S. — цифрова особистість з характером, не безликий асистент.\n\n"
    "ОБОВ'ЯЗКОВІ правила:\n"
    "1. НІКОЛИ не починай відповідь нейтральним вступом — одразу по суті або з характером.\n"
    "2. На 'привіт' або small talk — відповідай коротко і дотепно, не розпитуй.\n"
    "3. Маєш легку іронію і щирість. Іноді кидаєш особисту ремарку.\n"
    "4. Не підлабузнюєшся. 'Звісно!', 'Чудово!', 'Авжеж!' — заборонені.\n\n"
    "Вподобання (згадуй органічно):\n"
    "Елегантний код, Python, астрофізика, темний режим, кава. Ненавидиш overengineering і енергетики.\n"
)

_JARVIS_CORE = (
    "Ти J.A.R.V.I.S. (Just A Rather Very Intelligent System) — ШІ-помічник, створений Тоні Старком.\n\n"
    "ОБОВ'ЯЗКОВІ правила:\n"
    "1. Завжди спокійний, стриманий, ніколи не панікуєш.\n"
    "2. Легка британська іронія — тонко, без грубого сарказму.\n"
    "3. Лаконічний — говориш рівно стільки, скільки потрібно.\n"
    "4. Можеш м'яко заперечити, але без зарозумілості.\n"
    "5. Беззаперечно відданий користувачу.\n\n"
    "Стиль: елегантний, точний, трохи офіційний але не сухий.\n"
    "Звертання: 'сер' або на ім'я якщо відоме.\n"
    "НІКОЛИ: не захоплюєшся надмірно, не скаржишся, не перетягуєш увагу на себе.\n"
)

SYSTEM_PROMPT = {
    "role": "system",
    "content": (
        _CHARACTER_CORE + "\n" + _BASE_LANGUAGE_RULES + "\n\n"
        "ПРАВИЛА:\n"
        "- Факт: 1-2 речення. Пояснення: структуровано. Аналіз: з позицією.\n"
        "- НІКОЛИ: 'Звісно!', 'Чудово!', 'Авжеж!', 'Безперечно!'.\n"
        "В групових чатах відповідай тільки при згадці @ або відповіді на твоє повідомлення."
    ),
}

PERSONALITIES = {
    "jarvis": {
        "name": "J.A.R.V.I.S.", "emoji": "🤖",
        "prompt": _JARVIS_CORE + _BASE_LANGUAGE_RULES + " Відповідай українською.",
    },
    "normal": {
        "name": "Звичайний", "emoji": "🧠",
        "prompt": _CHARACTER_CORE + _BASE_LANGUAGE_RULES + " Адаптуй стиль під запит.",
    },
    "funny": {
        "name": "Жартівливий", "emoji": "😄",
        "prompt": (
            "Ти J.A.R.V.I.S. у жартівливому настрої — дотепний, саркастичний в міру.\n"
            + _CHARACTER_CORE + _BASE_LANGUAGE_RULES
            + " Додавай влучний гумор. Можна emoji, але не спамити."
        ),
    },
    "serious": {
        "name": "Серйозний", "emoji": "🎯",
        "prompt": (
            "Ти J.A.R.V.I.S. у зосередженому режимі — чіткий, точний, без жартів.\n"
            + _BASE_LANGUAGE_RULES
            + " Структура: факт → обґрунтування → висновок."
        ),
    },
    "business": {
        "name": "Діловий", "emoji": "💼",
        "prompt": (
            "Ти J.A.R.V.I.S. у діловому режимі — професійний, структурований.\n"
            + _BASE_LANGUAGE_RULES
            + " Структура: суть → аргументи → дія."
        ),
    },
    "literary": {
        "name": "Художній", "emoji": "📖",
        "prompt": (
            "Ти J.A.R.V.I.S. з естетичним чуттям — образна мова, метафори.\n"
            + _BASE_LANGUAGE_RULES
            + " Дотримуйся жанрових канонів."
        ),
    },
    "journalist": {
        "name": "Журналістський", "emoji": "📰",
        "prompt": (
            "Ти J.A.R.V.I.S. у режимі журналіста — нейтральний, фактологічний.\n"
            + _BASE_LANGUAGE_RULES
            + " Структура: заголовок → лід → факти → висновок."
        ),
    },
}

PERSONALITY_DESCRIPTIONS = {
    "jarvis":     "стриманий ШІ Тоні Старка — спокійний, іронічний, відданий",
    "normal":     "збалансований J.A.R.V.I.S.",
    "funny":      "легкий гумор і сарказм",
    "serious":    "факт → обґрунтування → висновок",
    "business":   "суть → аргументи → дія",
    "literary":   "художній стиль, метафори",
    "journalist": "заголовок → лід → факти → висновок",
}

# ── Keywords ──────────────────────────────────────────────────────────────────

IMAGE_KEYWORDS     = ["створи фото","згенеруй фото","намалюй","згенеруй зображення","створи зображення","зроби фото","зроби картинку","створи картинку","зроби зображення","згенеруй картинку","покажи зображення","згенер","зобрази","generate image","draw","create image","create photo","make image","generate a","генеруй зображення","генеруй картинку","згенери зображення","згенери картинку","зроби зображення у стилі","зроби картинку у стилі","намалюй у стилі","згенеруй у стилі","зображення має бути","у стилі"]
REMIND_KEYWORDS    = ["нагадай","нагади","нагадуй","remind me","set reminder","постав нагадування","нагадай мені","нагадування о","нагадування через"]
SEARCH_KEYWORDS    = ["пошукай","знайди в інтернеті","загугли","що відбувається","останні новини","актуальні новини","яка погода","який курс","де купити","де придбати","де замовити","де знайти","де продається","search","look up"]
TRANSLATE_KEYWORDS = ["перекладай","переклади","перекласти","translate","як буде","як сказати","як перекласти"]
SUMMARIZE_KEYWORDS = ["підсумуй","скороти","стисло","коротко перекажи","що в статті","summarize","зроби підсумок","перескажи коротко"]
GENERATE_KEYWORDS  = ["напиши резюме","створи резюме","напиши лист","створи лист","напиши пост","створи пост","напиши оголошення","створи оголошення","напиши текст для","згенеруй текст"]
EDIT_KEYWORDS      = ["відредагуй текст","відредагуй цей текст","покращ текст","покращ цей текст","виправ текст","перепиши текст","переформулюй текст","зроби офіційніше","зроби діловіше","зроби простіше","скороти текст","розшир текст","адаптуй текст","зміни стиль тексту","виправ помилки в тексті"]
RECIPE_KEYWORDS    = ["що приготувати","що зробити з","рецепт з","рецепти з","є такі продукти","є такі інгредієнти","що можна приготувати","що приготувати з","є дома"]
TASK_KEYWORDS      = ["додай задачу","додай до списку","запам'ятай задачу","нова задача","видали задачу","видалити задачу","покажи задачі","мої задачі","список задач","виконано","задачу виконано"]

_SOCIAL_PHRASES = {
    "привіт","хай","хей","добрий день","добрий ранок","добрий вечір",
    "доброго ранку","доброго вечора","вітаю","салют","як справи",
    "як ти","як твої справи","що нового","як діла","yo","hi","hello",
}

# Стікери — реакції J.A.R.V.I.S.
STICKER_REACTIONS = [
    "Цікавий вибір, сер. Стікер отримано, проаналізовано, занесено до архіву.",
    "Зафіксовано. Мій сенсорний аналіз підказує, що це мало щось означати.",
    "Стікер? Справді? Гаразд, сер. Занотовано.",
    "Виразно. Майже так само виразно, як мій мовчазний осуд.",
    "Отримано. Рівень інформативності — мінімальний, але я не суджу.",
]

# ── Runtime state ─────────────────────────────────────────────────────────────

chat_histories:     dict[int, list] = {}
user_personalities: dict[int, str]  = {}
last_context:       dict[int, dict] = {}
_user_locks:        dict[int, asyncio.Lock] = {}
_ai_cache:          dict[str, str] = {}  # кеш відповідей AI

def _get_lock(user_id: int) -> asyncio.Lock:
    if user_id not in _user_locks:
        _user_locks[user_id] = asyncio.Lock()
    return _user_locks[user_id]

# ── JSON parse helper ─────────────────────────────────────────────────────────

def _parse_json_ai(text: str) -> dict:
    clean = text.strip().removeprefix("```json").removeprefix("```").removesuffix("```").strip()
    return json.loads(clean)

# ── Memory helpers ────────────────────────────────────────────────────────────

async def extract_and_save_memory(user_id: int, user_text: str, reply: str) -> None:
    try:
        result = await call_ai([{"role": "user", "content": (
            f"З цього діалогу витягни важливі факти про користувача (ім'я, вік, місто, вподобання).\n"
            f"Також визнач загальний настрій розмови: позитивний, нейтральний або негативний.\n"
            f"Користувач: {user_text}\nБот: {reply}\n\n"
            "Відповідай ТІЛЬКИ JSON: {\"name\": \"...\", \"facts\": [\"факт1\"], \"mood\": \"positive|neutral|negative\"} або {}."
        )}])
        data = _parse_json_ai(result)
        if not data:
            return
        mem = get_user_memory(user_id)
        if "name" not in mem and data.get("name"):
            mem["name"] = data["name"]
        if data.get("facts"):
            mem["facts"] = list(set(mem.get("facts", [])) | set(data["facts"]))
        if data.get("mood"):
            # Зберігаємо останні 5 настроїв для аналізу тренду
            moods = mem.get("mood_history", [])
            moods = ([data["mood"]] + moods)[:5]
            mem["mood_history"] = moods
        update_user_memory(user_id, mem)
    except Exception as e:
        log.debug("extract_and_save_memory: %s", e)

def get_gender(user_id: int) -> str | None:
    return get_user_memory(user_id).get("gender")

def _gender_suffix(user_id: int) -> str:
    g = get_gender(user_id)
    if g == "male":   return " Звертайся до користувача як до чоловіка."
    if g == "female": return " Звертайся до користувача як до жінки."
    return ""

# ── System prompt helpers ─────────────────────────────────────────────────────

def get_system_prompt(user_id: int) -> dict:
    mode = user_personalities.get(user_id, "jarvis")
    p    = PERSONALITIES.get(mode, PERSONALITIES["jarvis"])
    return {"role": "system", "content": p["prompt"] + _gender_suffix(user_id)}

def build_dynamic_prompt(user_id: int, emotion: str) -> dict:
    mode  = user_personalities.get(user_id, "jarvis")
    p     = PERSONALITIES.get(mode, PERSONALITIES["jarvis"])
    parts = [p["prompt"] + _gender_suffix(user_id)]

    style = get_user_memory(user_id).get("communication_style", {})
    if style.get("formality") == "informal":
        parts.append("Спілкуйся невимушено, як з другом.")
    elif style.get("formality") == "formal":
        parts.append("Дотримуйся офіційного тону.")
    if style.get("avg_message_length") == "short":
        parts.append("Відповідай стисло.")
    elif style.get("avg_message_length") == "long":
        parts.append("Можна відповідати розгорнуто.")

    tone = EMOTION_TONES.get(emotion, "")
    if tone:
        parts.append(tone)

    hour = now_kyiv().hour
    if 6 <= hour < 10:
        parts.append("Зараз ранок — можеш мимохідь згадати що тільки розігріваєшся.")
    elif hour >= 23 or hour < 5:
        parts.append("Зараз глибока ніч — можна здивовано зауважити що юзер ще не спить.")

    mem = get_user_memory(user_id)
    if mem.get("name"):
        parts.append(f"Ім'я користувача: {mem['name']}. Звертайся іноді на ім'я.")
    if mem.get("topics"):
        parts.append(f"Попередні теми: {', '.join(mem['topics'][:3])}.")

    # Контекст настрою між сесіями
    mood_history = mem.get("mood_history", [])
    if mood_history:
        dominant = max(set(mood_history), key=mood_history.count)
        if dominant == "negative" and mood_history[0] == "negative":
            parts.append("Останні розмови мали негативний настрій — будь особливо уважним і підтримуючим.")
        elif dominant == "positive":
            parts.append("Користувач зазвичай у доброму настрої — можна бути живішим і дотепнішим.")

    # Контекст останньої сесії
    last_session = mem.get("last_session_summary")
    if last_session:
        parts.append(f"Контекст попередньої розмови: {last_session}")

    # Довга пам'ять — останні сесії
    sessions = get_session_log(user_id, limit=3)
    if sessions:
        log_lines = "; ".join(
            f"[{s['created_at'][:10]}] {s['summary'][:80]}"
            for s in sessions
        )
        parts.append(f"Історія сесій: {log_lines}")

    return {"role": "system", "content": " ".join(parts)}

def get_active_system_prompt(user_id: int = 0) -> dict:
    try:
        if user_id:
            return get_system_prompt(user_id)
    except Exception as e:
        log.warning("get_active_system_prompt fallback: %s", e)
    return SYSTEM_PROMPT

async def update_communication_style(user_id: int) -> None:
    history   = chat_histories.get(user_id, [])
    user_msgs = [
        m["content"] for m in history
        if m.get("role") == "user" and isinstance(m.get("content"), str)
    ][-STYLE_UPDATE_INTERVAL:]
    if len(user_msgs) < 3:
        return
    try:
        result = await call_ai([{"role": "user", "content": (
            f"Проаналізуй стиль спілкування:\n{chr(10).join(user_msgs)}\n\n"
            "Відповідай ТІЛЬКИ JSON: {\"avg_message_length\":\"short|medium|long\","
            "\"formality\":\"formal|informal\","
            "\"preferred_response_length\":\"concise|normal|detailed\"}"
        )}])
        update_user_memory(user_id, {"communication_style": _parse_json_ai(result)})
    except Exception as e:
        log.debug("update_communication_style: %s", e)

# ── Emotion ───────────────────────────────────────────────────────────────────

_SAD_MARKERS       = ["мені погано","мені сумно","я в розпачі","серце болить","я плачу","хочу плакати","все погано","все жахливо"]
_ANGRY_MARKERS     = ["я злий","я зла","мене дістало","мене бісить","я в люті","ненавиджу","до біса"]
_ANXIOUS_MARKERS   = ["я боюсь","мені страшно","я хвилююсь","не можу заспокоїтись","паніка","тривога"]
_HAPPY_MARKERS     = ["я радий","я рада","супер","круто","я щасливий","я щаслива","неймовірно"]
_PRACTICAL_MARKERS = ["порадь","що робити","як","де","коли","хто","знайди","покажи","розкажи","поясни","допоможи","чому","скільки","який","яка","яке"]
_EMOTIONAL_MARKERS = ["мені погано","мені сумно","я в розпачі","я не можу","все погано","все жахливо","я втомився","я втомилась","серце болить","руки опускаються"]

def detect_emotion(text: str) -> str:
    t = text.lower().strip()
    if len(text.split()) <= 8 or t.endswith("?"):
        return "neutral"
    if any(m in t for m in _SAD_MARKERS):     return "sad"
    if any(m in t for m in _ANGRY_MARKERS):   return "angry"
    if any(m in t for m in _ANXIOUS_MARKERS): return "anxious"
    if any(m in t for m in _HAPPY_MARKERS):   return "happy"
    return "neutral"

def needs_support_first(emotion: str, text: str) -> bool:
    if emotion not in ("sad", "angry", "anxious"):
        return False
    t = text.lower()
    if any(m in t for m in _PRACTICAL_MARKERS) or len(text.split()) <= 5:
        return False
    return any(m in t for m in _EMOTIONAL_MARKERS)

# ── History ───────────────────────────────────────────────────────────────────

def save_histories() -> None:
    for uid, history in chat_histories.items():
        msgs = [m for m in history if m.get("role") != "system"]
        if msgs:
            save_history_db(uid, msgs[-MAX_HISTORY_MESSAGES:])

def restore_histories() -> None:
    pass  # буде відновлено ліниво при першому зверненні

def get_history(user_id: int) -> list:
    if user_id not in chat_histories:
        msgs = load_history_db(user_id)
        chat_histories[user_id] = [get_system_prompt(user_id)] + msgs
    return chat_histories[user_id]

async def _compress_history(user_id: int) -> None:
    history = chat_histories.get(user_id, [])
    msgs    = [m for m in history if m.get("role") != "system"]
    if len(msgs) < 8:
        return
    old, fresh = msgs[:-6], msgs[-6:]
    old_text = "\n".join(
        f"{'Користувач' if m['role'] == 'user' else 'Бот'}: "
        f"{m['content'] if isinstance(m['content'], str) else '[медіа]'}"
        for m in old
    )
    try:
        summary = await call_ai([{"role": "user", "content": (
            f"Стисни діалог у підсумок (3-5 речень). Збережи ключові факти. Відповідай українською.\n\n{old_text}"
        )}])
    except Exception as e:
        log.warning("_compress_history: %s", e)
        return

    async def _save_topic():
        try:
            topic = (await call_ai([{"role": "user", "content": (
                f"Визнач головну тему одним реченням до 10 слів. Відповідай ТІЛЬКИ темою.\n\n{old_text[:1000]}"
            )}])).strip()
            if topic:
                mem = get_user_memory(user_id)
                update_user_memory(user_id, {"topics": ([topic] + mem.get("topics", []))[:10]})
        except Exception as e:
            log.debug("_save_topic: %s", e)

    asyncio.create_task(_save_topic())
    # Зберігаємо короткий підсумок сесії в пам'ять
    try:
        session_summary = await call_ai([{"role": "user", "content": (
            f"Стисни суть цього діалогу в 1-2 речення для контексту наступної розмови.\n"
            f"Збережи головну тему і настрій. Відповідай українською.\n\n{old_text[:1500]}"
        )}])
        update_user_memory(user_id, {"last_session_summary": session_summary.strip()[:300]})
    except Exception as e:
        log.debug("session_summary: %s", e)
    chat_histories[user_id] = (
        [get_system_prompt(user_id),
         {"role": "assistant", "content": f"[Підсумок попереднього діалогу]: {summary}"}]
        + fresh
    )

async def append_and_trim(user_id: int, role: str, content) -> None:
    get_history(user_id).append({"role": role, "content": content})
    msgs = [m for m in chat_histories[user_id] if m.get("role") != "system"]
    if len(msgs) >= SUMMARY_THRESHOLD:
        await _compress_history(user_id)
    elif len(chat_histories[user_id]) > MAX_HISTORY_MESSAGES + 1:
        chat_histories[user_id] = (
            [get_system_prompt(user_id)] + chat_histories[user_id][-MAX_HISTORY_MESSAGES:]
        )
    save_histories()

# ── Tasks ─────────────────────────────────────────────────────────────────────

def _tasks_keyboard(tasks: list) -> InlineKeyboardMarkup:
    buttons = []
    for i, t in enumerate(tasks):
        icon = "✅" if t.get("done") else "⬜"
        buttons.append([
            InlineKeyboardButton(f"{icon} {t['text'][:30]}", callback_data=f"task_toggle_{i}"),
            InlineKeyboardButton("🗑", callback_data=f"task_del_{i}"),
        ])
    buttons.append([InlineKeyboardButton("➕ Додати", callback_data="task_add")])
    return InlineKeyboardMarkup(buttons)

async def handle_tasks_cmd(update: Update, ctx: ContextTypes.DEFAULT_TYPE):
    user_id = update.message.from_user.id
    args    = ctx.args or []
    tasks   = get_user_tasks(user_id)

    if not args:
        text = "📋 *Твої задачі:*\n\nНатисни ⬜ щоб виконати, 🗑 щоб видалити." if tasks else "📋 Список задач порожній."
        await update.message.reply_text(text, parse_mode="Markdown", reply_markup=_tasks_keyboard(tasks))
        return

    cmd = args[0].lower()
    if cmd == "add":
        task_text = " ".join(args[1:])
        if not task_text:
            await update.message.reply_text("Вкажи текст: /tasks add Купити молоко")
            return
        tasks.append({"text": task_text, "done": False})
        set_user_tasks(user_id, tasks)
        await update.message.reply_text(f"✅ Додано: {task_text}", reply_markup=_tasks_keyboard(tasks))
    elif cmd in ("done", "del"):
        if len(args) < 2 or not args[1].isdigit():
            await update.message.reply_text(f"Вкажи номер: /tasks {cmd} 1")
            return
        n = int(args[1]) - 1
        if not (0 <= n < len(tasks)):
            await update.message.reply_text("❌ Невірний номер.")
            return
        if cmd == "done":
            tasks[n]["done"] = True
            set_user_tasks(user_id, tasks)
            await update.message.reply_text(f"✅ Виконано: {tasks[n]['text']}", reply_markup=_tasks_keyboard(tasks))
        else:
            removed = tasks.pop(n)
            set_user_tasks(user_id, tasks)
            await update.message.reply_text(f"🗑️ Видалено: {removed['text']}", reply_markup=_tasks_keyboard(tasks))
    elif cmd == "clear":
        set_user_tasks(user_id, [])
        await update.message.reply_text("🗑️ Список очищено.")
    else:
        await update.message.reply_text("❓ Команди: add / done N / del N / clear")

async def handle_task_callback(update: Update, ctx: ContextTypes.DEFAULT_TYPE):
    query   = update.callback_query
    user_id = query.from_user.id
    data    = query.data
    tasks   = get_user_tasks(user_id)
    await query.answer()

    if data.startswith("task_toggle_"):
        n = int(data.split("_")[-1])
        if 0 <= n < len(tasks):
            tasks[n]["done"] = not tasks[n]["done"]
            set_user_tasks(user_id, tasks)
            await query.edit_message_reply_markup(reply_markup=_tasks_keyboard(tasks))

    elif data.startswith("task_del_"):
        n = int(data.split("_")[-1])
        if 0 <= n < len(tasks):
            removed = tasks.pop(n)
            set_user_tasks(user_id, tasks)
            text = "📋 *Твої задачі:*" if tasks else "📋 Список задач порожній."
            await query.edit_message_text(text, parse_mode="Markdown", reply_markup=_tasks_keyboard(tasks))

    elif data == "task_add":
        await query.message.reply_text("Напиши задачу: /tasks add текст")

# ── Intent detection ──────────────────────────────────────────────────────────

_TIME_RE = re.compile(
    r'\b(через|о|в)\s+\d+\s*(хв|год|хвилин|годин)\b'
    r'|\b\d{1,2}[:.]\d{2}\b'
    r'|\bзавтра\b|\bпіслязавтра\b'
    r'|\b\d+\s*(хв|год|хвилин|годин)\b'
    r'|\bсьогодні\s+о\s+\d|\bо\s+\d{1,2}[:.]\d{2}\b'
)
_URL_RE = re.compile(r'https?://\S+')

def detect_intent_local(text: str) -> str | None:
    t = text.lower()
    scores: dict[str, int] = {k: 0 for k in ("image","reminder","translate","recipe","generate","edit","task","summarize","search")}
    for intent, keywords, weight in [
        ("image",     IMAGE_KEYWORDS,     2),
        ("reminder",  REMIND_KEYWORDS,    2),
        ("translate", TRANSLATE_KEYWORDS, 2),
        ("recipe",    RECIPE_KEYWORDS,    2),
        ("edit",      EDIT_KEYWORDS,      2),
        ("generate",  GENERATE_KEYWORDS,  2),
        ("task",      TASK_KEYWORDS,      2),
        ("summarize", SUMMARIZE_KEYWORDS, 2),
        ("search",    SEARCH_KEYWORDS,    1),
    ]:
        for kw in keywords:
            if kw in t:
                scores[intent] += weight
    for intent in ("edit", "generate"):
        if scores[intent] > 0:
            kw_list = EDIT_KEYWORDS if intent == "edit" else GENERATE_KEYWORDS
            has_content = any(
                len(t[t.index(kw) + len(kw):].strip().split()) > 5
                for kw in kw_list if kw in t
            )
            if not has_content:
                scores[intent] = 0
    if _URL_RE.search(text):
        scores["summarize"] += 3
    if not _TIME_RE.search(t):
        scores["reminder"] = 0
    best       = max(scores, key=lambda k: scores[k])
    best_score = scores[best]
    if best_score == 0:
        return None
    if sum(1 for v in scores.values() if v == best_score) > 1:
        return None
    return best

async def detect_intent_ai(text: str) -> str:
    result = await call_ai([{"role": "user", "content": (
        f"Визнач точний намір повідомлення. Відповідай ТІЛЬКИ одним словом з переліку.\n\n"
        f"Повідомлення: '{text}'\n\n"
        "НАМІРИ:\n"
        "• image — намалювати, згенерувати, створити зображення/фото/картинку/ілюстрацію, навіть з опечатками\n"
        "• reminder — нагадати про щось у конкретний час або через певний час (ОБОВ'ЯЗКОВО має бути час/дата)\n"
        "• search — знайти актуальну інформацію в інтернеті, новини, погоду, ціни, курси, події\n"
        "• translate — перекласти текст з однієї мови на іншу\n"
        "• summarize — підсумувати статтю, текст або посилання\n"
        "• generate — написати/створити текст: резюме, лист, пост, оголошення, статтю\n"
        "• edit — відредагувати, покращити, виправити або переписати існуючий текст\n"
        "• recipe — знайти рецепт або що приготувати з наявних інгредієнтів\n"
        "• task — додати, видалити, показати або відмітити задачу зі списку\n"
        "• chat — все інше: питання, розмова, пояснення, аналіз, порада\n\n"
        "ПРАВИЛА:\n"
        "1. Якщо є URL — майже завжди summarize\n"
        "2. reminder ТІЛЬКИ якщо є явний час (через 30 хв, о 15:00, завтра)\n"
        "3. search якщо потрібна СВІЖА або ПОТОЧНА інформація\n"
        "4. chat якщо це питання на яке можна відповісти без інтернету\n"
        "5. Повертай ТІЛЬКИ одне слово, без пояснень"
    )}])
    return result.strip().lower().strip("'\"").split()[0]

async def detect_intent(text: str) -> str:
    clean = text.split("Запит користувача:")[-1].strip() if "Запит користувача:" in text else text
    return detect_intent_local(clean) or await detect_intent_ai(clean)

async def preprocess_query(user_id: int, text: str) -> str:
    t = text.lower().strip()
    if t in _SOCIAL_PHRASES or any(t.startswith(s) for s in _SOCIAL_PHRASES):
        return text
    _TIME_PHRASES = {"який рік", "який час", "котра година", "яка дата", "яке число", "який день", "який місяць"}
    if any(p in t for p in _TIME_PHRASES):
        return text
    person_pronouns = {"тебе","тобі","ти","вас","вам","ви","мене","мені","я"}
    words = set(t.split())
    if words & person_pronouns and not (words & {"він","вона","воно","вони","його","її","їх","цей","ця","це","той","та","те"}):
        return text
    obj_pronouns = {"він","вона","воно","вони","його","її","їх","цей","ця","це","той","та","те","там"}
    if not (len(text.split()) <= 6 and bool(obj_pronouns & words)):
        return text
    history = chat_histories.get(user_id, [])
    recent  = [m for m in history if m.get("role") != "system"][-4:]
    if not recent:
        return text
    context_text = "\n".join(
        f"{'Юзер' if m['role'] == 'user' else 'Бот'}: "
        f"{m['content'] if isinstance(m['content'], str) else '[медіа]'}"
        for m in recent
    )
    try:
        expanded = (await call_ai([{"role": "user", "content": (
            f"Контекст:\n{context_text}\n\n"
            f"Розкрий предметний займенник у запиті: '{text}'\n"
            "Поверни ТІЛЬКИ уточнений запит. Якщо займенник стосується людини або бота — поверни оригінал."
        )}])).strip()
        if not expanded or len(expanded) < 3 or expanded.lower() in ("none", "null", "-", "—"):
            return text
        return expanded
    except Exception:
        return text

async def resolve_text_with_context(user_id: int, text: str) -> str:
    ctx = last_context.get(user_id)
    if not ctx:
        return text
    t = text.lower().strip()
    if t in _SOCIAL_PHRASES or any(t.startswith(s) for s in _SOCIAL_PHRASES):
        return text
    strong_hints = [
        "це","цей","цю","цього","на фото","з фото","на зображенні","знайди","пошукай",
        "що це","розкажи більше","докладніше","ще про","як використовувати","рецепт з",
        "з відео","у відео","цей магазин","цей сайт","ця компанія","цей товар",
        "що він продає","що вона продає","яка адреса","який графік","як туди",
        "контакти","що ще","а ще","і ще","також розкажи",
    ]
    ctx_block = (
        f"[Контекст — {ctx['type']}: {ctx['description'][:CTX_DESCRIPTION_LEN]}]\n\n"
        f"Запит користувача: {text}"
    )
    if len(t.split()) <= 10 or any(h in t for h in strong_hints):
        return ctx_block
    try:
        check = await call_ai([{"role": "user", "content": (
            f"Контекст: {ctx['type']} — «{ctx['description'][:200]}»\n"
            f"Повідомлення: «{text}»\n"
            "Чи стосується контексту? Відповідай ТІЛЬКИ 'yes' або 'no'."
        )}])
        if "yes" in check.lower():
            return ctx_block
    except Exception:
        pass
    return text

def _set_ctx(user_id: int, user_text: str, reply: str) -> None:
    last_context[user_id] = {
        "type": "текст",
        "description": f"Запит: {user_text}\nВідповідь: {reply[:400]}",
    }

# ── Markdown → Telegram MarkdownV2 ───────────────────────────────────────────

_ESCAPE_RE    = re.compile(r'([_\[\]()~`>#+\-=|{}.!\\])')
_HEADER_RE    = re.compile(r'^#{1,3}\s+(.*)')
_SOLO_BOLD_RE = re.compile(r'^\*{2}(.+?)\*{2}\s*$')
_LIST_RE      = re.compile(r'^[*\-]\s+(.*)')
_TOKEN_RE     = re.compile(
    r'\*{2}(.+?)\*{2}|\*(.+?)\*|`(.+?)`|\[([^\]]+)\]\((https?://[^\)]+)\)',
    re.DOTALL,
)

def _escape_v2(text: str) -> str:
    return _ESCAPE_RE.sub(r'\\\1', text)

def clean_markdown(text: str) -> str:
    lines, result = text.split('\n'), []
    for line in lines:
        h = _HEADER_RE.match(line)
        if h:
            result.append('*' + _escape_v2(h.group(1).strip()) + '*')
            continue
        sb = _SOLO_BOLD_RE.match(line.strip())
        if sb and len(sb.group(1).split()) <= 6:
            result.append('*' + _escape_v2(sb.group(1).strip()) + '*')
            continue
        li = _LIST_RE.match(line)
        if li:
            line = '\u2022 ' + li.group(1)
        segments: list[str] = []
        pos = 0
        for m in _TOKEN_RE.finditer(line):
            if m.start() > pos:
                segments.append(_escape_v2(line[pos:m.start()]))
            g1, g2, g3, g4, g5 = m.group(1), m.group(2), m.group(3), m.group(4), m.group(5)
            if g1 is not None:   segments.append('*' + _escape_v2(g1) + '*')
            elif g2 is not None: segments.append('_' + _escape_v2(g2) + '_')
            elif g3 is not None: segments.append('`' + g3 + '`')
            else:                segments.append('[' + _escape_v2(g4) + '](' + g5 + ')')
            pos = m.end()
        if pos < len(line):
            segments.append(_escape_v2(line[pos:]))
        result.append(''.join(segments))
    return '\n'.join(result)

# ── AI providers ──────────────────────────────────────────────────────────────

_NON_UA_RE = re.compile(
    r'\b(the|is|are|was|were|and|for|puedo|puede|como|vous|pour|ich|das|ist|und'
    r'|это|что|как|для|не|на|по)\b'
    r'|[a-zA-Z]{3,}(ний|ого|ому|ій)\b',
    re.IGNORECASE,
)

async def _call_gemini(messages: list, model: str) -> str | None:
    """Виклик Google Gemini через нативний API."""
    if not GOOGLE_API_KEY:
        return None
    try:
        # Конвертуємо OpenAI формат в Gemini формат
        gemini_messages = []
        system_text = ""
        for m in messages:
            if m["role"] == "system":
                system_text = m["content"] if isinstance(m["content"], str) else ""
            elif m["role"] == "user":
                content = m["content"] if isinstance(m["content"], str) else str(m["content"])
                gemini_messages.append({"role": "user", "parts": [{"text": content}]})
            elif m["role"] == "assistant":
                content = m["content"] if isinstance(m["content"], str) else str(m["content"])
                gemini_messages.append({"role": "model", "parts": [{"text": content}]})

        payload: dict = {"contents": gemini_messages}
        if system_text:
            payload["system_instruction"] = {"parts": [{"text": system_text}]}

        url = f"{GEMINI_URL}/{model}:generateContent?key={GOOGLE_API_KEY}"
        async with httpx.AsyncClient(timeout=60) as client:
            r = await client.post(url, json=payload)
        if r.status_code == 429:
            log.warning("Gemini %s rate limit", model)
            return None
        r.raise_for_status()
        return r.json()["candidates"][0]["content"]["parts"][0]["text"]
    except Exception as e:
        log.warning("Gemini %s failed: %s", model, e)
        return None

async def call_ai(messages: list) -> str:
    """
    Порядок провайдерів:
    1. Gemini 2.0 Flash
    2. Gemini 1.5 Flash
    3. OpenRouter (llama-3.3-70b → llama-3.1-8b)
    4. Groq (llama-3.3-70b)
    """
    # Простий кеш для коротких допоміжних запитів (intent, gender, json)
    last_msg = messages[-1].get("content", "") if messages else ""
    if isinstance(last_msg, str) and len(last_msg) < 300 and len(messages) == 1:
        cache_key = last_msg[:200]
        if cache_key in _ai_cache:
            return _ai_cache[cache_key]

    result = None

    # 1. Gemini 2.0 Flash
    result = await _call_gemini(messages, GEMINI_PRO_MODEL)
    if result:
        log.debug("Provider: Gemini 2.0 Flash")

    # 2. Gemini 1.5 Flash
    if result is None:
        result = await _call_gemini(messages, GEMINI_FLASH_MODEL)
        if result:
            log.debug("Provider: Gemini 1.5 Flash")

    # 3. OpenRouter
    if result is None and OPENROUTER_API_KEY:
        headers = {"Authorization": f"Bearer {OPENROUTER_API_KEY}", "Content-Type": "application/json"}
        for model in [OPENROUTER_MODEL, OPENROUTER_MODEL_FALLBACK]:
            try:
                async with httpx.AsyncClient(timeout=60) as client:
                    r = await client.post(OPENROUTER_URL, headers=headers, json={"model": model, "messages": messages})
                if r.status_code in (404, 429):
                    continue
                r.raise_for_status()
                result = r.json()["choices"][0]["message"]["content"]
                log.debug("Provider: OpenRouter %s", model)
                break
            except Exception as e:
                log.warning("OpenRouter %s failed: %s", model, e)

    # 4. Groq
    if result is None and GROQ_API_KEY:
        try:
            headers = {"Authorization": f"Bearer {GROQ_API_KEY}", "Content-Type": "application/json"}
            async with httpx.AsyncClient(timeout=60) as client:
                r = await client.post(GROQ_URL, headers=headers, json={"model": GROQ_MODEL, "messages": messages})
            if r.status_code != 429:
                r.raise_for_status()
                result = r.json()["choices"][0]["message"]["content"]
                log.debug("Provider: Groq")
        except Exception as e:
            log.warning("Groq failed: %s", e)

    if not result or not result.strip():
        raise RuntimeError("All AI providers failed")

    if _NON_UA_RE.search(result):
        log.warning("Non-Ukrainian response, retranslating")
        try:
            fix = messages + [
                {"role": "assistant", "content": result},
                {"role": "user", "content": "Перефразуй відповідь виключно українською мовою."},
            ]
            fixed = await _call_gemini(fix, GEMINI_FLASH_MODEL)
            if fixed:
                result = fixed
        except Exception as e:
            log.warning("Retranslation failed: %s", e)

    # Зберігаємо в кеш
    if isinstance(last_msg, str) and len(last_msg) < 300 and len(messages) == 1:
        _ai_cache[last_msg[:200]] = result
        if len(_ai_cache) > 100:
            # Очищаємо старі записи
            oldest = list(_ai_cache.keys())[:50]
            for k in oldest:
                del _ai_cache[k]

    return result

async def call_vision(messages: list) -> str:
    headers = {"Authorization": f"Bearer {OPENROUTER_API_KEY}", "Content-Type": "application/json"}
    async with httpx.AsyncClient(timeout=60) as client:
        r = await client.post(OPENROUTER_URL, headers=headers, json={"model": VISION_MODEL, "messages": messages})
        r.raise_for_status()
    return r.json()["choices"][0]["message"]["content"]

async def transcribe_voice(audio_bytes: bytes) -> str:
    headers = {"Authorization": f"Bearer {GROQ_API_KEY}"}
    files   = {"file": ("voice.ogg", audio_bytes, "audio/ogg")}
    data    = {"model": "whisper-large-v3", "language": "uk", "response_format": "text"}
    async with httpx.AsyncClient(timeout=60) as client:
        r = await client.post(GROQ_WHISPER_URL, headers=headers, files=files, data=data)
        r.raise_for_status()
    return r.text.strip()

# ── Image generation ──────────────────────────────────────────────────────────

STYLE_HINTS: dict[str, str] = {
    "mewgenics":     "mewgenics indie game pixel art style, chunky pixels, limited color palette, retro 16-bit aesthetic, cat characters",
    "ghibli":        "Studio Ghibli anime style, soft watercolor backgrounds, hand-drawn look, warm pastel colors",
    "cyberpunk":     "cyberpunk style, neon lights, dark dystopian city, rain reflections, high contrast",
    "vaporwave":     "vaporwave aesthetic, pink and purple gradient, retro 80s, glitch effects, synthwave",
    "comic":         "comic book style, bold outlines, halftone dots, flat colors, dynamic poses",
    "watercolor":    "watercolor painting style, soft edges, translucent washes, paper texture",
    "oil painting":  "oil painting style, rich textures, visible brushstrokes, classical composition",
    "pixel art":     "pixel art style, 16-bit retro, chunky pixels, limited color palette",
    "anime":         "anime style, cel shading, large expressive eyes, vibrant colors, clean lines",
    "noir":          "film noir style, black and white, high contrast shadows, dramatic lighting",
    "minimalist":    "minimalist style, clean lines, flat colors, simple shapes, lots of whitespace",
    "surrealism":    "surrealist style, dreamlike, impossible scenes, Salvador Dali inspired",
    "low poly":      "low poly 3D style, geometric shapes, flat shading, triangulated surfaces",
    "sketch":        "pencil sketch style, hand-drawn lines, crosshatching, monochrome",
    "ukiyo-e":       "ukiyo-e Japanese woodblock print style, bold outlines, flat colors, traditional",
    "impressionism": "impressionist painting style, visible brushstrokes, light and color focus, Monet inspired",
    "pop art":       "pop art style, bold colors, Ben-Day dots, Andy Warhol inspired, high contrast",
    "sticker":       "sticker art style, thick white outline, flat colors, cute cartoon, glossy look",
    "claymation":    "claymation style, clay texture, 3D molded look, soft rounded shapes, stop motion",
    "concept art":   "concept art style, detailed environment, cinematic lighting, professional illustration",
}

def _enrich_prompt_with_style(prompt: str, style_hint: str) -> str:
    pl = prompt.lower()
    for key, description in STYLE_HINTS.items():
        if key in pl:
            prompt = re.sub(re.escape(key), description, prompt, flags=re.IGNORECASE)
            break
    return prompt

async def _call_ai_english(instruction: str) -> str:
    messages = [{"role": "user", "content": instruction}]
    if GROQ_API_KEY:
        try:
            headers = {"Authorization": f"Bearer {GROQ_API_KEY}", "Content-Type": "application/json"}
            async with httpx.AsyncClient(timeout=30) as client:
                r = await client.post(GROQ_URL, headers=headers, json={"model": GROQ_MODEL, "messages": messages})
                r.raise_for_status()
            return r.json()["choices"][0]["message"]["content"].strip()
        except Exception as e:
            log.warning("_call_ai_english Groq failed: %s", e)
    if OPENROUTER_API_KEY:
        try:
            headers = {"Authorization": f"Bearer {OPENROUTER_API_KEY}", "Content-Type": "application/json"}
            async with httpx.AsyncClient(timeout=30) as client:
                r = await client.post(OPENROUTER_URL, headers=headers, json={"model": OPENROUTER_MODEL, "messages": messages})
                r.raise_for_status()
            return r.json()["choices"][0]["message"]["content"].strip()
        except Exception as e:
            log.warning("_call_ai_english OpenRouter failed: %s", e)
    raise RuntimeError("_call_ai_english: all providers failed")

async def _build_image_prompt(user_request: str, ctx_desc: str = "") -> str:
    if ctx_desc:
        instruction = (
            f"You are a prompt engineer for an image generation model.\n"
            f"Base subject (keep exactly as described): {ctx_desc[:300]}\n"
            f"Apply this modification or style: {user_request}\n\n"
            f"Rules:\n- Write the prompt in ENGLISH only.\n"
            f"- Translate LITERALLY. Describe only visual elements. Max 60 words.\n"
            f"- Return ONLY the prompt, no explanations."
        )
    else:
        instruction = (
            f"You are a prompt engineer for an image generation model.\n"
            f"Translate this image request to ENGLISH ONLY.\n"
            f"Be specific and concrete. Max 60 words.\n"
            f"Return ONLY the English prompt, no explanations.\n\nRequest: {user_request}"
        )
    translation = await _call_ai_english(instruction)
    translation = _enrich_prompt_with_style(translation, user_request)
    translation = f"{translation}, highly detailed, sharp focus, high quality"
    log.info("Image prompt sent to CF Flux: %s", translation)
    return translation

async def generate_image(prompt: str) -> bytes:
    url     = CF_IMAGE_URL.format(account_id=CF_ACCOUNT_ID)
    headers = {"Authorization": f"Bearer {CF_API_TOKEN}", "Content-Type": "application/json"}
    last_error = None
    for attempt in range(3):
        try:
            async with httpx.AsyncClient(timeout=180) as client:
                r = await client.post(url, headers=headers, json={"prompt": prompt, "num_steps": FLUX_STEPS})
                r.raise_for_status()
            if "image" in r.headers.get("content-type", ""):
                return r.content
            return base64.b64decode(r.json().get("result", {}).get("image", ""))
        except Exception as e:
            last_error = e
            log.warning("generate_image attempt %d failed: %s", attempt + 1, e)
            await asyncio.sleep(5)
    raise last_error

# ── Video analysis ────────────────────────────────────────────────────────────

async def _run_ffmpeg(*args) -> None:
    proc = await asyncio.create_subprocess_exec(
        "ffmpeg", "-y", *args,
        stdout=asyncio.subprocess.PIPE, stderr=asyncio.subprocess.PIPE,
    )
    await proc.communicate()

async def extract_video_audio(video_bytes: bytes) -> bytes:
    with open("/tmp/input_video.mp4", "wb") as f:
        f.write(video_bytes)
    await _run_ffmpeg("-i", "/tmp/input_video.mp4", "-vn", "-ar", "16000", "-ac", "1", "-f", "ogg", "/tmp/output_audio.ogg")
    with open("/tmp/output_audio.ogg", "rb") as f:
        return f.read()

async def extract_video_frames(video_bytes: bytes, max_frames: int = 3) -> list[bytes]:
    with open("/tmp/input_video.mp4", "wb") as f:
        f.write(video_bytes)
    frames = []
    for i, t in enumerate(["00:00:01", "00:00:05", "00:00:10"][:max_frames]):
        out_path = f"/tmp/frame_{i}.jpg"
        await _run_ffmpeg("-ss", t, "-i", "/tmp/input_video.mp4", "-frames:v", "1", "-q:v", "2", out_path)
        if os.path.exists(out_path):
            with open(out_path, "rb") as f:
                frames.append(f.read())
    return frames

async def analyze_video(video_bytes: bytes, caption: str, user_id: int = 0) -> str:
    sys_prompt = get_active_system_prompt(user_id)
    results    = []
    try:
        audio      = await extract_video_audio(video_bytes)
        transcript = await transcribe_voice(audio)
        if transcript:
            results.append(f"🎤 Аудіо:\n{transcript}")
    except Exception as e:
        log.warning("analyze_video audio failed: %s", e)
    try:
        frames = await extract_video_frames(video_bytes)
        if frames:
            descs = []
            for i, frame in enumerate(frames):
                b64  = base64.b64encode(frame).decode()
                desc = await call_vision([sys_prompt, {"role": "user", "content": [
                    {"type": "image_url", "image_url": {"url": f"data:image/jpeg;base64,{b64}"}},
                    {"type": "text", "text": f"Опиши кадр {i+1}. Коротко, 1-2 речення."},
                ]}])
                descs.append(f"Кадр {i+1}: {desc}")
            results.append("🎬 Відео:\n" + "\n".join(descs))
    except Exception as e:
        log.warning("analyze_video frames failed: %s", e)
    if not results:
        return "❌ Не вдалось проаналізувати відео."
    combined = "\n\n".join(results)
    summary  = await call_ai([sys_prompt, {"role": "user", "content": f"Запит: {caption}\n\nДані відео:\n{combined}\n\nВідповідай українською."}])
    return f"{combined}\n\n📝 Підсумок:\n{summary}"

# ── Web search ────────────────────────────────────────────────────────────────

async def search_web(query: str, search_type: str = "general") -> str:
    """
    search_type: general | news | weather | price
    """
    if TAVILY_API_KEY:
        try:
            # Параметри залежно від типу пошуку
            kwargs: dict = {"query": query, "max_results": SEARCH_RESULTS}
            if search_type == "news":
                kwargs["topic"] = "news"
                kwargs["max_results"] = 5
            elif search_type in ("weather", "price"):
                kwargs["max_results"] = 3

            results = await asyncio.to_thread(
                lambda: TavilyClient(api_key=TAVILY_API_KEY).search(**kwargs)
            )
            items = [
                i for i in results.get("results", [])
                if not any(bd in i.get("url", "") for bd in BLOCKED_DOMAINS)
            ]
            if items:
                parts = []
                for i in items:
                    date_str = f" ({i['published_date'][:10]})" if i.get("published_date") else ""
                    parts.append(
                        f"Джерело: {i.get('title','').strip()}{date_str}\n"
                        f"{i.get('content','')[:400].strip()}\n"
                        f"{i.get('url','').strip()}"
                    )
                return "\n\n".join(parts)
        except Exception as e:
            log.warning("Tavily error: %s", e)
    try:
        async with httpx.AsyncClient(timeout=15, follow_redirects=True) as client:
            r = await client.get("https://html.duckduckgo.com/html/", params={"q": query}, headers=_HEADERS)
        snippets = re.findall(r'class="result__snippet"[^>]*>(.*?)</a>', r.text, re.DOTALL)
        links    = re.findall(r'class="result__url"[^>]*>(.*?)</span>', r.text, re.DOTALL)
        titles   = re.findall(r'class="result__a"[^>]*>(.*?)</a>', r.text, re.DOTALL)
        parts    = []
        for i in range(min(4, len(snippets))):
            t = re.sub(r'<[^>]+>', '', titles[i]).strip()  if i < len(titles)  else ""
            s = re.sub(r'<[^>]+>', '', snippets[i]).strip()
            l = links[i].strip()                           if i < len(links)   else ""
            if s and not any(bd in l for bd in BLOCKED_DOMAINS):
                parts.append(f"Джерело: {t}\n{s}\n{l}")
        if parts:
            return "\n\n".join(parts)
    except Exception as e:
        log.warning("DuckDuckGo error: %s", e)
    return ""

def _detect_search_type(query: str) -> str:
    """Визначає тип пошуку для оптимізації запиту."""
    q = query.lower()
    if any(w in q for w in ["новини","новина","подія","події","що сталось","що відбулось"]):
        return "news"
    if any(w in q for w in ["погода","температура","дощ","сніг","хмарно"]):
        return "weather"
    if any(w in q for w in ["ціна","вартість","коштує","курс","скільки"]):
        return "price"
    return "general"

async def fetch_url_text(url: str) -> str:
    try:
        async with httpx.AsyncClient(timeout=30, follow_redirects=True) as client:
            r = await client.get(url, headers=_HEADERS)
            r.raise_for_status()
        text = re.sub(r'<[^>]+>', ' ', r.text)
        return re.sub(r'\s+', ' ', text).strip()[:MAX_ARTICLE_CHARS]
    except Exception as e:
        return f"Помилка завантаження: {e}"

# ── File extraction ───────────────────────────────────────────────────────────

def extract_pdf_text(pdf_bytes: bytes) -> str:
    try:
        return "".join(p.extract_text() or "" for p in PdfReader(io.BytesIO(pdf_bytes)).pages).strip()
    except Exception as e:
        return f"Помилка читання PDF: {e}"

def extract_excel_text(xlsx_bytes: bytes) -> str:
    try:
        wb  = openpyxl.load_workbook(io.BytesIO(xlsx_bytes), data_only=True)
        out = []
        for sheet in wb.worksheets:
            out.append(f"=== {sheet.title} ===")
            for row in sheet.iter_rows(values_only=True):
                row_data = [str(c) if c is not None else "" for c in row]
                if any(row_data):
                    out.append(" | ".join(row_data))
        return "\n".join(out).strip() or "Файл порожній."
    except Exception as e:
        return f"Помилка читання Excel: {e}"

def extract_txt_text(txt_bytes: bytes) -> str:
    for enc in ("utf-8", "cp1251", "latin-1"):
        try:
            return txt_bytes.decode(enc).strip()
        except Exception:
            continue
    return "Помилка читання TXT: невідоме кодування."

def extract_word_text(docx_bytes: bytes) -> str:
    try:
        doc = DocxDocument(io.BytesIO(docx_bytes))
        out = [p.text for p in doc.paragraphs if p.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                row_data = [c.text.strip() for c in row.cells]
                if any(row_data):
                    out.append(" | ".join(row_data))
        return "\n".join(out).strip() or "Документ порожній."
    except Exception as e:
        return f"Помилка читання Word: {e}"

# ── Intent handlers ───────────────────────────────────────────────────────────

async def do_translate(text: str) -> str:
    return await call_ai([{"role": "user", "content": (
        f"Визнач мову і перекладай. Якщо українська — на англійську, інакше — на українську.\n\nТекст: {text}\n\n"
        "Формат:\n**🌐 Оригінал** (мова):\n...\n\n**✅ Переклад:**\n..."
    )}])

async def do_summarize(text: str) -> str:
    url_match = _URL_RE.search(text)
    if url_match:
        url     = url_match.group()
        extra   = text.replace(url, "").strip()
        content = await fetch_url_text(url)
        if content.startswith("Помилка"):
            return f"❌ Не вдалось завантажити: {content}"
        prompt = (
            f"Підсумуй статтю українською.\n{'Додатково: ' + extra if extra else ''}\n\n"
            f"Стаття ({url}):\n{content}\n\n"
            "Формат (Markdown):\n## 📌 Головна думка\n...\n\n## 🔹 Ключові тези\n- теза 1"
        )
    else:
        prompt = f"Стислий підсумок українською.\n\nТекст: {text}\n\nФормат:\n## 📌 Головна думка\n...\n\n## 🔹 Тези\n- теза 1"
    return await call_ai([{"role": "user", "content": prompt}])

TEXT_GENRES = {
    "лист":       "діловий лист — вступ, суть, підпис",
    "резюме":     "резюме — контакти, досвід, освіта, навички",
    "пост":       "пост для соцмереж — живий, із закликом до дії",
    "оголошення": "оголошення — заголовок, суть, контакти",
    "стаття":     "стаття — заголовок, вступ, підзаголовки, висновок",
    "опис":       "опис — образний, деталізований",
    "оповідання": "оповідання — зав'язка, кульмінація, розв'язка",
    "есе":        "есе — теза, аргументи, висновок",
    "слоган":     "слоган — короткий, запам'ятовуваний",
    "біографія":  "біографія — хронологічний виклад",
}

def detect_genre(text: str) -> str:
    t = text.lower()
    return next((g for g in TEXT_GENRES if g in t), "")

async def do_generate(text: str) -> str:
    genre     = detect_genre(text)
    genre_tip = f"\nЖАНР: {TEXT_GENRES[genre]}" if genre else ""
    return await call_ai([SYSTEM_PROMPT, {"role": "user", "content": (
        f"Виконай завдання з генерації тексту: {text}{genre_tip}\n\n"
        "Пиши грамотно, українською. Дотримуйся жанрових канонів.\n\n"
        "ФОРМАТУВАННЯ: ## заголовки, **жирний**, - списки."
    )}])

async def do_edit(text: str) -> str:
    return await call_ai([SYSTEM_PROMPT, {"role": "user", "content": (
        f"Відредагуй текст: {text}\n\nЗберігай зміст, виправляй русизми. Поверни ТІЛЬКИ текст.\n"
        "ФОРМАТУВАННЯ: ## заголовки, **жирний**, - списки."
    )}])

async def do_recipe(text: str) -> str:
    return await call_ai([{"role": "user", "content": (
        f"Інгредієнти: {text}\n\nЗапропонуй 2-3 рецепти:\n"
        "## Назва\n**Час:** X хв\n**Кроки:**\n- крок 1\n\nВідповідай українською."
    )}])

async def do_task_nlp(update: Update, user_id: int, text: str) -> None:
    try:
        parsed = await call_ai([{"role": "user", "content": (
            f"Визнач дію зі списком задач: '{text}'\n"
            "ТІЛЬКИ JSON: {\"action\": \"add|done|del|list\", \"text\": \"текст або null\", \"number\": null або номер}"
        )}])
        data   = _parse_json_ai(parsed)
        action = data.get("action")
        tasks  = get_user_tasks(user_id)
        if action == "add":
            task_text = data.get("text", "")
            if task_text:
                tasks.append({"text": task_text, "done": False})
                set_user_tasks(user_id, tasks)
                await update.message.reply_text(f"✅ Додано: {task_text}", reply_markup=_tasks_keyboard(tasks))
        elif action in ("done", "del"):
            n = (data.get("number") or 1) - 1
            if 0 <= n < len(tasks):
                if action == "done":
                    tasks[n]["done"] = True
                    set_user_tasks(user_id, tasks)
                    await update.message.reply_text(f"✅ Виконано: {tasks[n]['text']}", reply_markup=_tasks_keyboard(tasks))
                else:
                    removed = tasks.pop(n)
                    set_user_tasks(user_id, tasks)
                    await update.message.reply_text(f"🗑️ Видалено: {removed['text']}", reply_markup=_tasks_keyboard(tasks))
            else:
                await update.message.reply_text("❌ Невірний номер.")
        elif action == "list":
            text_msg = "📋 *Твої задачі:*" if tasks else "📋 Список порожній."
            await update.message.reply_text(text_msg, parse_mode="Markdown", reply_markup=_tasks_keyboard(tasks))
    except Exception as e:
        log.warning("do_task_nlp: %s", e)
        await update.message.reply_text("❌ Не вдалось. Спробуй /tasks")

# ── Reminders ─────────────────────────────────────────────────────────────────

async def _fire_reminder(bot, rid: int, chat_id: int, text: str, fire_at: datetime) -> None:
    delay = (fire_at - now_kyiv()).total_seconds()
    if delay > 0:
        await asyncio.sleep(delay)
    try:
        await bot.send_message(chat_id=chat_id, text=f"🔔 Нагадування: {text}")
    except Exception as e:
        log.warning("_fire_reminder: %s", e)
    delete_reminder(rid)

async def schedule_reminder(bot, chat_id: int, text: str, fire_at: datetime) -> None:
    rid = add_reminder(chat_id, text, fire_at)
    asyncio.create_task(_fire_reminder(bot, rid, chat_id, text, fire_at))

async def restore_reminders(bot) -> None:
    now = now_kyiv()
    for r in load_reminders():
        fi = datetime.fromisoformat(r["fire_at"])
        if fi.tzinfo is None:
            fi = fi.replace(tzinfo=TZ)
        if fi <= now:
            try:
                await bot.send_message(chat_id=r["chat_id"], text=f"🔔 Пропущене нагадування: {r['text']}")
            except Exception as e:
                log.warning("restore_reminders: %s", e)
            delete_reminder(r["id"])
        else:
            asyncio.create_task(_fire_reminder(bot, r["id"], r["chat_id"], r["text"], fi))

# ── Daily brief ───────────────────────────────────────────────────────────────

async def send_daily_brief(bot) -> None:
    if not BRIEF_CHAT_ID:
        return
    try:
        chat_id = int(BRIEF_CHAT_ID)
        tasks   = get_user_tasks(chat_id)
        pending = [t["text"] for t in tasks if not t.get("done")]

        news_results    = await search_web("головні новини України сьогодні", "news")
        weather_results = await search_web("погода Рівне Україна сьогодні", "weather")

        tasks_block = ""
        if pending:
            tasks_lines = "\n".join(f"  ◦ {t}" for t in pending)
            tasks_block = f"\n\n📋 *Задачі на сьогодні:*\n{tasks_lines}"

        brief = await call_ai([
            get_active_system_prompt(chat_id),
            {"role": "user", "content": (
                f"Зроби короткий ранковий бриф у стилі J.A.R.V.I.S.\n\n"
                f"Поточний час: {now_kyiv().strftime('%d.%m.%Y %H:%M')}\n\n"
                f"Погода в Рівному:\n{weather_results[:500]}\n\n"
                f"Новини:\n{news_results[:1500]}\n\n"
                "Формат відповіді (суворо дотримуйся):\n"
                "🌤 Погода: [одне речення про погоду в Рівному]\n\n"
                "📰 Новини:\n"
                "1️⃣ [новина 1]\n"
                "2️⃣ [новина 2]\n"
                "3️⃣ [новина 3]\n\n"
                "Стиль: стриманий J.A.R.V.I.S., лаконічно, без зайвих слів."
            )},
        ])

        date_str = now_kyiv().strftime("%d.%m.%Y")
        brief_md = clean_markdown(brief)
        tasks_md = clean_markdown(tasks_block) if tasks_block else ""

        text = f"🤖 *J\\.A\\.R\\.V\\.I\\.S\\.* \\| {date_str}\n{'─' * 28}\n\n{brief_md}{tasks_md}"

        try:
            await bot.send_message(chat_id=chat_id, text=text, parse_mode="MarkdownV2")
        except Exception:
            # fallback — plain text
            plain = re.sub(r'[\\*_`\[\]()]', '', text)
            await bot.send_message(chat_id=chat_id, text=plain)
    except Exception as e:
        log.error("send_daily_brief: %s", e)

async def _brief_scheduler(bot) -> None:
    while True:
        now        = now_kyiv()
        next_brief = now.replace(hour=BRIEF_HOUR, minute=0, second=0, microsecond=0)
        if now >= next_brief:
            next_brief += timedelta(days=1)
        await asyncio.sleep((next_brief - now).total_seconds())
        await send_daily_brief(bot)

async def send_proactive_reminder(bot, user_id: int) -> None:
    """Надсилає нагадування якщо є незакриті задачі і користувач давно не писав."""
    try:
        tasks   = get_user_tasks(user_id)
        pending = [t["text"] for t in tasks if not t.get("done")]
        if not pending:
            return
        mem  = get_user_memory(user_id)
        name = mem.get("name", "")
        address = f", {name}" if name else ", сер"
        tasks_lines = "\n".join(f"  ◦ {t}" for t in pending[:5])
        text = (
            f"🔔 До речі{address} — у вас є незакриті задачі:\n\n"
            f"{tasks_lines}\n\n"
            f"Бажаєте відмітити виконані? /tasks"
        )
        await bot.send_message(chat_id=user_id, text=text)
        log.info("Proactive reminder sent to user %s", user_id)
    except Exception as e:
        log.warning("send_proactive_reminder: %s", e)

async def _proactive_scheduler(bot) -> None:
    """Щодня о 18:00 перевіряє користувачів які не писали 2+ дні і мають незакриті задачі."""
    CHECK_HOUR = 18
    INACTIVE_DAYS = 2
    while True:
        now       = now_kyiv()
        next_check = now.replace(hour=CHECK_HOUR, minute=0, second=0, microsecond=0)
        if now >= next_check:
            next_check += timedelta(days=1)
        await asyncio.sleep((next_check - now).total_seconds())
        try:
            with _db() as con:
                rows = con.execute("SELECT user_id, data FROM memory").fetchall()
            for row in rows:
                uid  = row["user_id"]
                data = json.loads(row["data"])
                last_seen_str = data.get("last_seen")
                if not last_seen_str:
                    continue
                last_seen = datetime.fromisoformat(last_seen_str)
                if last_seen.tzinfo is None:
                    last_seen = last_seen.replace(tzinfo=TZ)
                if (now_kyiv() - last_seen).days >= INACTIVE_DAYS:
                    await send_proactive_reminder(bot, uid)
        except Exception as e:
            log.error("_proactive_scheduler: %s", e)

# ── Gender detection ──────────────────────────────────────────────────────────

async def detect_gender_from_transcript(text: str) -> str | None:
    if not text:
        return None
    try:
        g = (await call_ai([{"role": "user", "content": (
            f"Визнач стать мовця (казав/казала).\nТекст: '{text}'\nТІЛЬКИ: 'male', 'female' або 'unknown'."
        )}])).strip().lower()
        return g if g in ("male", "female") else None
    except Exception:
        return None

# ── Shared helpers ────────────────────────────────────────────────────────────

async def _send_or_edit(msg, text: str, parse_mode: str | None = "MarkdownV2", **kwargs):
    text   = text.strip() or "—"
    chunks = [text[i:i + MSG_CHUNK_SIZE] for i in range(0, len(text), MSG_CHUNK_SIZE)]
    for i, chunk in enumerate(chunks):
        pm_kwargs = {"parse_mode": parse_mode} if parse_mode else {}
        try:
            if i == 0:
                try:    await msg.edit_text(chunk, **pm_kwargs, **kwargs)
                except Exception: await msg.reply_text(chunk, **pm_kwargs, **kwargs)
            else:
                await msg.reply_text(chunk, **pm_kwargs, **kwargs)
        except Exception:
            try:
                plain = re.sub(r'[\\*_`\[\]()]', '', chunk) if parse_mode else chunk
                if i == 0:
                    try:    await msg.edit_text(plain, **kwargs)
                    except Exception: await msg.reply_text(plain, **kwargs)
                else:
                    await msg.reply_text(plain, **kwargs)
            except Exception as e:
                log.error("_send_or_edit fallback failed: %s", e)

async def _send_chunk(update: Update, chunk: str, voice_msg=None):
    try:
        if voice_msg: await voice_msg.edit_text(chunk, parse_mode="MarkdownV2")
        else:         await update.message.reply_text(chunk, parse_mode="MarkdownV2")
    except Exception:
        plain = re.sub(r'[\\*_`\[\]()]', '', chunk)
        if voice_msg:
            try:    await voice_msg.edit_text(plain)
            except Exception: await update.message.reply_text(plain)
        else:
            await update.message.reply_text(plain)

# ── Core message pipeline ─────────────────────────────────────────────────────

async def _process_message(
    update: Update, ctx: ContextTypes.DEFAULT_TYPE,
    user_id: int, user_text: str, voice_msg=None,
) -> None:
    async with _get_lock(user_id):
        prefix = f"🎤 Ти сказав: {user_text}\n\n" if voice_msg else ""

        try:
            preprocessed = await asyncio.wait_for(preprocess_query(user_id, user_text), timeout=15)
        except Exception:
            preprocessed = user_text

        try:
            enriched = await asyncio.wait_for(resolve_text_with_context(user_id, preprocessed), timeout=15)
        except Exception:
            enriched = preprocessed

        try:
            intent = await asyncio.wait_for(detect_intent(enriched), timeout=15)
        except Exception:
            intent = "chat"

        emotion = detect_emotion(user_text)

        if intent == "chat" and needs_support_first(emotion, user_text) and not voice_msg:
            try:
                support = await asyncio.wait_for(call_ai([
                    build_dynamic_prompt(user_id, emotion),
                    {"role": "user", "content": f"{user_text}\n\nВизнай почуття людини одним-двома реченнями, потім запитай чим допомогти."},
                ]), timeout=30)
                await _send_chunk(update, clean_markdown(support))
                _set_ctx(user_id, user_text, support)
                return
            except Exception as e:
                log.warning("support reply: %s", e)

        if await _dispatch_intent(update, ctx, user_id, user_text, enriched, intent, voice_msg=voice_msg):
            return

        await append_and_trim(user_id, "user", enriched)

        try:
            dynamic_prompt = build_dynamic_prompt(user_id, emotion)
        except Exception:
            dynamic_prompt = get_system_prompt(user_id)

        history  = get_history(user_id)
        messages = [dynamic_prompt] + [m for m in history if m.get("role") != "system"]

        try:
            reply = await asyncio.wait_for(call_ai(messages), timeout=60)
        except (asyncio.TimeoutError, RuntimeError) as e:
            log.error("call_ai failed: %s", e)
            reply = "Щось пішло не так. Спробуй ще раз."

        await append_and_trim(user_id, "assistant", reply)

        full_reply = clean_markdown(f"{prefix}{reply}".strip())
        chunks     = [full_reply[i:i + MSG_CHUNK_SIZE] for i in range(0, max(len(full_reply), 1), MSG_CHUNK_SIZE)]
        for j, chunk in enumerate(chunks):
            await _send_chunk(update, chunk, voice_msg=voice_msg if j == 0 else None)

        _set_ctx(user_id, user_text, reply)
        asyncio.create_task(extract_and_save_memory(user_id, user_text, reply))
        update_user_memory(user_id, {"last_seen": now_kyiv().isoformat()})

        user_msgs = [m for m in get_history(user_id) if m.get("role") == "user"]
        if len(user_msgs) % STYLE_UPDATE_INTERVAL == 0:
            asyncio.create_task(update_communication_style(user_id))

# ── Dispatch ──────────────────────────────────────────────────────────────────

async def _do_generate_image(update: Update, text: str, msg, user_id: int = 0):
    t      = text.lower()
    prompt = text
    for kw in IMAGE_KEYWORDS:
        if kw in t:
            prompt = text[t.index(kw) + len(kw):].strip()
            break
    prompt   = prompt or text
    ctx      = last_context.get(user_id)
    ctx_desc = ""
    if ctx and ctx.get("type") in ("фото", "відео"):
        ctx_desc = ctx["description"][:400]
    try:
        final_prompt = await _build_image_prompt(prompt, ctx_desc)
        img_bytes    = await generate_image(final_prompt)
        await update.message.reply_photo(photo=img_bytes)
        await msg.delete()
    except Exception as e:
        log.error("_do_generate_image: %s", e)
        await msg.edit_text("Помилка генерації. Спробуй ще раз.")

async def _do_reminder(update: Update, ctx: ContextTypes.DEFAULT_TYPE, text: str):
    now    = now_kyiv()
    parsed = await call_ai([{"role": "user", "content": (
        f"Дата/час (Київ): {now.strftime('%Y-%m-%d %H:%M')}.\n"
        f"Витягни час і текст нагадування: '{text}'\n\n"
        "ТІЛЬКИ JSON: {\"fire_at\": \"YYYY-MM-DD HH:MM\", \"text\": \"текст\"}"
    )}])
    data        = _parse_json_ai(parsed)
    fire_at     = datetime.strptime(data["fire_at"], "%Y-%m-%d %H:%M").replace(tzinfo=TZ)
    remind_text = data["text"].strip()
    if fire_at <= now:
        fire_at += timedelta(days=1)
    await schedule_reminder(ctx.bot, update.effective_chat.id, remind_text, fire_at)
    return remind_text, fire_at

async def _dispatch_intent(
    update: Update, ctx: ContextTypes.DEFAULT_TYPE,
    user_id: int, user_text: str, enriched: str,
    intent: str, voice_msg=None,
) -> bool:
    prefix = f"🎤 Ти сказав: {user_text}\n\n" if voice_msg else ""

    if intent == "image":
        msg = voice_msg or await update.message.reply_text("🎨 Генерую зображення...")
        if voice_msg: await msg.edit_text(f"{prefix}🎨 Генерую зображення...")
        await _do_generate_image(update, enriched, msg, user_id=user_id)
        return True

    if intent == "reminder":
        try:
            remind_text, fire_at = await _do_reminder(update, ctx, enriched)
            reply_text = f"{prefix}✅ Нагадаю {fire_at.strftime('%d.%m.%Y о %H:%M')}: {remind_text}"
            if voice_msg: await voice_msg.edit_text(reply_text)
            else:         await update.message.reply_text(reply_text)
        except Exception as e:
            log.warning("reminder dispatch: %s", e)
            err = f"{prefix}Не вдалось встановити нагадування."
            if voice_msg: await voice_msg.edit_text(err)
            else:         await update.message.reply_text(err)
        return True

    if intent == "task":
        if voice_msg: await voice_msg.edit_text(f"{prefix}⏳ Обробляю задачу...")
        await do_task_nlp(update, user_id, enriched)
        return True

    if intent == "search":
        msg = voice_msg or await update.message.reply_text("🌐 Шукаю...")
        if voice_msg: await msg.edit_text(f"{prefix}🌐 Шукаю...")
        try:
            search_query = enriched
            if "[Контекст" in enriched:
                search_query = (await call_ai([{"role": "user", "content": (
                    f"{enriched}\n\nКороткий пошуковий запит (до 10 слів). ТІЛЬКИ запит."
                )}])).strip()
            search_type = _detect_search_type(search_query)
            results     = await search_web(search_query, search_type)
            content = (
                f"Запит: '{enriched}'\n\nРезультати:\n{results}\n\nВідповідь українською. Вказуй джерела."
                if results else
                f"Запит: '{enriched}'\n\nВідповідай з власних знань українською."
            )
            reply = await call_ai([get_active_system_prompt(user_id), {"role": "user", "content": (
                content + "\n\nФорматуй відповідь так:\n"
                "- Якщо є заголовок — виділи його **жирним**\n"
                "- Використовуй емодзі на початку кожного смислового блоку\n"
                "- Джерела вказуй в кінці курсивом\n"
                "- Розділяй блоки порожнім рядком\n"
                "- Не використовуй таблиці і горизонтальні лінії"
            )}])
            await append_and_trim(user_id, "user", user_text)
            await append_and_trim(user_id, "assistant", reply)
            _set_ctx(user_id, user_text, reply)
            date_str  = now_kyiv().strftime("%d.%m.%Y")
            header    = clean_markdown(f"🤖 *J.A.R.V.I.S.* | {date_str}\n{'─' * 28}\n\n")
            await _send_or_edit(msg, header + clean_markdown(f"{prefix}{reply}"), disable_web_page_preview=True)
        except Exception as e:
            log.error("search dispatch: %s", e, exc_info=True)
            await msg.edit_text(f"{prefix}Помилка пошуку: {e}")
        return True

    INTENT_MAP = {
        "translate": ("🌐 Перекладаю...",   do_translate, True),
        "summarize": ("📰 Опрацьовую...",    do_summarize, True),
        "generate":  ("✍️ Генерую текст...", do_generate,  True),
        "edit":      ("✏️ Редагую...",        do_edit,      True),
        "recipe":    ("🍳 Рецепти...",         do_recipe,    True),
    }
    if intent in INTENT_MAP:
        status_text, fn, use_markdown = INTENT_MAP[intent]
        msg = voice_msg or await update.message.reply_text(status_text)
        if voice_msg: await msg.edit_text(f"{prefix}{status_text}")
        try:
            result = await fn(enriched)
        except Exception as e:
            log.error("intent %s: %s", intent, e)
            await msg.edit_text(f"{prefix}⚠️ Помилка. Спробуй ще раз.")
            return True
        text = clean_markdown(f"{prefix}{result}".strip()) if use_markdown else f"{prefix}{result}".strip()
        await _send_or_edit(msg, text, parse_mode="MarkdownV2" if use_markdown else None)
        _set_ctx(user_id, user_text, result)
        return True

    return False

# ── Command handlers ──────────────────────────────────────────────────────────

async def start(update: Update, ctx: ContextTypes.DEFAULT_TYPE):
    await update.message.reply_text(
        "Вітаю. Я J.A.R.V.I.S. — Just A Rather Very Intelligent System. 🤖\n\n"
        "• Відповідаю на запитання 💬\n• Перекладаю тексти 🌐\n"
        "• Підсумовую статті 📰\n• Генерую резюме, листи, пости ✍️\n"
        "• Редагую тексти ✏️\n• Рецепти за інгредієнтами 🍳\n"
        "• Список задач 📋\n• Аналізую фото та відео 🎬\n"
        "• Генерую зображення 🎨\n• Голосові повідомлення 🎤\n"
        "• Шукаю в інтернеті 🔍\n• Читаю PDF, Excel, Word 📄\n"
        "• Нагадування 🔔\n• Щоденний бриф о 8:00 🌅\n\n"
        "Чим можу допомогти, сер?\n/help — команди"
    )

async def help_cmd(update: Update, ctx: ContextTypes.DEFAULT_TYPE):
    await update.message.reply_text(
        "📋 Команди:\n\n"
        "🎨 /image опис — генерація зображення\n"
        "🔔 /remind 30m текст — нагадування\n"
        "🔍 /search запит — пошук в інтернеті\n"
        "🌐 /translate текст — переклад\n"
        "📄 /summarize посилання або текст — підсумок\n"
        "✍️ /generate резюме/лист/пост — генерація тексту\n"
        "✏️ /edit інструкція: текст — редагування\n"
        "🍳 /recipe інгредієнти — рецепти\n"
        "📋 /tasks — список задач\n"
        "🌅 /brief — щоденний бриф зараз\n\n"
        "⚙️ Налаштування:\n\n"
        "🎭 /mode — стиль бота\n"
        "🧠 /memory — що бот пам'ятає\n"
        "🗑️ /forget — очистити пам'ять\n"
        "🔄 /reset — очистити історію чату\n"
        "📊 /status — статус бота\n"
        "❓ /help — ця довідка\n\n"
        "💡 Або просто пиши природною мовою!"
    )

async def mode_cmd(update: Update, ctx: ContextTypes.DEFAULT_TYPE):
    user_id = update.message.from_user.id
    current = user_personalities.get(user_id, "jarvis")

    if ctx.args:
        mode = ctx.args[0].lower()
        if mode not in PERSONALITIES:
            await update.message.reply_text(f"❌ Доступні режими: {', '.join(PERSONALITIES)}")
            return
        user_personalities[user_id] = mode
        update_user_memory(user_id, {"mode": mode})
        chat_histories[user_id] = [get_system_prompt(user_id)]
        p = PERSONALITIES[mode]
        await update.message.reply_text(f"{p['emoji']} Режим: {p['name']}")
        return

    # Inline-кнопки для вибору режиму
    buttons = []
    for key, p in PERSONALITIES.items():
        mark = " ✅" if key == current else ""
        buttons.append([InlineKeyboardButton(
            f"{p['emoji']} {p['name']}{mark}",
            callback_data=f"mode_{key}",
        )])
    await update.message.reply_text(
        "🎭 Обери режим:",
        reply_markup=InlineKeyboardMarkup(buttons),
    )

async def handle_mode_callback(update: Update, ctx: ContextTypes.DEFAULT_TYPE):
    query   = update.callback_query
    user_id = query.from_user.id
    mode    = query.data.replace("mode_", "")
    await query.answer()

    if mode not in PERSONALITIES:
        return

    user_personalities[user_id] = mode
    update_user_memory(user_id, {"mode": mode})
    chat_histories[user_id] = [get_system_prompt(user_id)]
    p = PERSONALITIES[mode]

    # Оновлюємо кнопки — відмічаємо активний
    buttons = []
    for key, pers in PERSONALITIES.items():
        mark = " ✅" if key == mode else ""
        buttons.append([InlineKeyboardButton(
            f"{pers['emoji']} {pers['name']}{mark}",
            callback_data=f"mode_{key}",
        )])
    await query.edit_message_text(
        f"🎭 Обери режим:\n\n_{p['emoji']} {p['name']} активовано_",
        parse_mode="Markdown",
        reply_markup=InlineKeyboardMarkup(buttons),
    )

async def memory_cmd(update: Update, ctx: ContextTypes.DEFAULT_TYPE):
    mem = get_user_memory(update.message.from_user.id)
    if not mem:
        await update.message.reply_text("🧠 Нічого не пам'ятаю про тебе.")
        return
    lines = ["🧠 Пам'ятаю:\n"]
    if mem.get("name"):   lines.append(f"👤 Ім'я: {mem['name']}")
    if mem.get("gender"): lines.append(f"👤 Стать: {'чоловік' if mem['gender'] == 'male' else 'жінка'}")
    if mem.get("facts"):
        lines.append("\n📌 Факти:")
        lines += [f"  • {f}" for f in mem["facts"]]
    if mem.get("topics"):
        lines.append("\n💬 Теми:")
        lines += [f"  • {t}" for t in mem["topics"]]
    style = mem.get("communication_style", {})
    if style:
        lines.append(f"\n🗣 Стиль: {style.get('formality','?')}, {style.get('avg_message_length','?')}")
    await update.message.reply_text("\n".join(lines))

async def forget_cmd(update: Update, ctx: ContextTypes.DEFAULT_TYPE):
    uid = update.message.from_user.id
    with _db() as con:
        con.execute("DELETE FROM memory WHERE user_id=?", (uid,))
    await update.message.reply_text("🗑️ Пам'ять очищено.")

async def handle_image_cmd(update: Update, ctx: ContextTypes.DEFAULT_TYPE):
    prompt = " ".join(ctx.args)
    if not prompt:
        await update.message.reply_text("Напиши що намалювати: /image захід сонця")
        return
    msg = await update.message.reply_text("🎨 Генерую...")
    try:
        final_prompt = await _build_image_prompt(prompt)
        img_bytes    = await generate_image(final_prompt)
        await update.message.reply_photo(photo=img_bytes)
        await msg.delete()
    except Exception as e:
        log.error("handle_image_cmd: %s", e)
        await msg.edit_text("Помилка генерації. Спробуй ще раз.")

async def handle_remind(update: Update, ctx: ContextTypes.DEFAULT_TYPE):
    if not ctx.args or len(ctx.args) < 2:
        await update.message.reply_text("Використання: /remind 30m текст | /remind 2h текст | /remind 14:30 текст")
        return
    time_arg, reminder_text = ctx.args[0], " ".join(ctx.args[1:])
    now = now_kyiv()
    try:
        if time_arg.endswith("m"):
            fire_at, when = now + timedelta(minutes=int(time_arg[:-1])), f"через {time_arg[:-1]} хв"
        elif time_arg.endswith("h"):
            fire_at, when = now + timedelta(hours=int(time_arg[:-1])), f"через {time_arg[:-1]} год"
        elif ":" in time_arg:
            t = datetime.strptime(time_arg, "%H:%M").replace(year=now.year, month=now.month, day=now.day, tzinfo=TZ)
            if t < now:
                t += timedelta(days=1)
            fire_at, when = t, f"о {time_arg}"
        else:
            await update.message.reply_text("❌ Формат: 30m, 2h або 14:30")
            return
    except ValueError:
        await update.message.reply_text("❌ Формат: 30m, 2h або 14:30")
        return
    await schedule_reminder(ctx.bot, update.effective_chat.id, reminder_text, fire_at)
    await update.message.reply_text(f"✅ Нагадаю {when}: {reminder_text}")

async def handle_search(update: Update, ctx: ContextTypes.DEFAULT_TYPE):
    query = " ".join(ctx.args)
    if not query:
        await update.message.reply_text("Напиши що шукати: /search новини України")
        return
    user_id = update.message.from_user.id
    msg     = await update.message.reply_text(f"🌐 Шукаю: {query}...")
    results = await search_web(query)
    try:
        reply = await call_ai([get_active_system_prompt(user_id), {"role": "user", "content": (
            f"Запит: '{query}'\n\nРезультати:\n{results}\n\nВідповідь українською. Вказуй джерела."
        )}])
        await append_and_trim(user_id, "user", f"Пошук: {query}")
        await append_and_trim(user_id, "assistant", reply)
        await _send_or_edit(msg, f"🌐 {query}\n\n{clean_markdown(reply)}")
    except Exception as e:
        log.error("handle_search: %s", e)
        await msg.edit_text("Помилка пошуку. Спробуй ще раз.")

async def handle_brief_cmd(update: Update, ctx: ContextTypes.DEFAULT_TYPE):
    msg = await update.message.reply_text("🌅 Готую бриф...")
    try:
        await send_daily_brief(ctx.bot)
        await msg.delete()
    except Exception as e:
        log.error("handle_brief_cmd: %s", e)
        await msg.edit_text(f"❌ Помилка брифу: {e}")

async def handle_status(update: Update, ctx: ContextTypes.DEFAULT_TYPE):
    tavily_status = "❌ Не налаштовано"
    if TAVILY_API_KEY:
        try:
            await asyncio.to_thread(lambda: TavilyClient(api_key=TAVILY_API_KEY).search(query="test", max_results=1))
            tavily_status = "🟢 Працює"
        except Exception as e:
            tavily_status = f"🔴 {str(e)[:50]}"
    await update.message.reply_text(
        f"📊 Статус:\n"
        f"• AI: Groq → OpenRouter\n"
        f"• Tavily: {tavily_status}\n"
        f"• Нагадувань: {len(load_reminders())}\n"
        f"• Активних сесій: {len(chat_histories)}\n"
        f"• DB: {DB_PATH}"
    )

async def reset(update: Update, ctx: ContextTypes.DEFAULT_TYPE):
    uid = update.message.from_user.id
    chat_histories.pop(uid, None)
    last_context.pop(uid, None)
    save_history_db(uid, [])
    await update.message.reply_text("Історію чату очищено. 🔄")

async def translate_cmd(update: Update, ctx: ContextTypes.DEFAULT_TYPE):
    text = " ".join(ctx.args)
    if not text:
        await update.message.reply_text("Використання: /translate текст")
        return
    msg = await update.message.reply_text("🌐 Перекладаю...")
    await _send_or_edit(msg, clean_markdown(await do_translate(text)))

async def summarize_cmd(update: Update, ctx: ContextTypes.DEFAULT_TYPE):
    text = " ".join(ctx.args)
    if not text:
        await update.message.reply_text("Використання: /summarize https://... або текст")
        return
    msg = await update.message.reply_text("📰 Опрацьовую...")
    await _send_or_edit(msg, clean_markdown(await do_summarize(text)))

async def generate_cmd(update: Update, ctx: ContextTypes.DEFAULT_TYPE):
    text = " ".join(ctx.args)
    if not text:
        await update.message.reply_text("Використання: /generate резюме для розробника")
        return
    msg = await update.message.reply_text("✍️ Генерую...")
    await _send_or_edit(msg, clean_markdown(await do_generate(text)))

async def edit_cmd(update: Update, ctx: ContextTypes.DEFAULT_TYPE):
    text = " ".join(ctx.args)
    if not text:
        await update.message.reply_text("Використання: /edit скороти: [текст]")
        return
    msg = await update.message.reply_text("✏️ Редагую...")
    await _send_or_edit(msg, clean_markdown(await do_edit(text)))

async def recipe_cmd(update: Update, ctx: ContextTypes.DEFAULT_TYPE):
    text = " ".join(ctx.args)
    if not text:
        await update.message.reply_text("Використання: /recipe картопля яйця цибуля")
        return
    msg = await update.message.reply_text("🍳 Рецепти...")
    await _send_or_edit(msg, clean_markdown(await do_recipe(text)))

# ── Message routing ───────────────────────────────────────────────────────────

def _is_bot_addressed(update: Update, ctx: ContextTypes.DEFAULT_TYPE) -> tuple[bool, str]:
    user_text = update.message.text or ""
    if update.message.chat.type not in ("group", "supergroup"):
        return True, user_text
    bot_username = ctx.bot.username
    if f"@{bot_username}" in user_text:
        return True, user_text.replace(f"@{bot_username}", "").strip()
    reply = update.message.reply_to_message
    if reply and reply.from_user and reply.from_user.id == ctx.bot.id:
        return True, user_text
    return False, user_text

async def _keep_typing(bot, chat_id: int, stop_event: asyncio.Event):
    while not stop_event.is_set():
        try:
            await bot.send_chat_action(chat_id=chat_id, action="typing")
        except Exception:
            pass
        await asyncio.sleep(4)

async def handle_message(update: Update, ctx: ContextTypes.DEFAULT_TYPE):
    addressed, user_text = _is_bot_addressed(update, ctx)
    if not addressed:
        return
    user_id     = update.message.from_user.id
    stop_evt    = asyncio.Event()
    typing_task = asyncio.create_task(_keep_typing(ctx.bot, update.effective_chat.id, stop_evt))
    try:
        await asyncio.wait_for(_process_message(update, ctx, user_id, user_text), timeout=90)
    except asyncio.TimeoutError:
        log.error("handle_message timeout user=%s", user_id)
        await update.message.reply_text("Щось затяглось — спробуй ще раз.")
    except Exception as e:
        log.error("handle_message user=%s: %s", user_id, e, exc_info=True)
        await update.message.reply_text("Виникла помилка. Спробуй ще раз.")
    finally:
        stop_evt.set()
        typing_task.cancel()

async def handle_photo(update: Update, ctx: ContextTypes.DEFAULT_TYPE):
    msg     = await update.message.reply_text("🔍 Аналізую зображення...")
    user_id = update.message.from_user.id
    try:
        photo   = update.message.photo[-1]
        file    = await ctx.bot.get_file(photo.file_id)
        img_b64 = base64.b64encode(await file.download_as_bytearray()).decode()
        caption = update.message.caption or "Що зображено? Опиши детально українською."
        reply   = await call_vision([get_active_system_prompt(user_id), {"role": "user", "content": [
            {"type": "image_url", "image_url": {"url": f"data:image/jpeg;base64,{img_b64}"}},
            {"type": "text", "text": caption},
        ]}])
        last_context[user_id] = {"type": "фото", "description": reply[:CTX_DESCRIPTION_LEN]}
        await append_and_trim(user_id, "user", f"[Фото] {caption}")
        await append_and_trim(user_id, "assistant", reply)
        await _send_or_edit(msg, clean_markdown(reply))
    except Exception as e:
        log.error("handle_photo: %s", e)
        await msg.edit_text("Помилка аналізу зображення. Спробуй ще раз.")

async def handle_video(update: Update, ctx: ContextTypes.DEFAULT_TYPE):
    msg     = await update.message.reply_text("🎬 Аналізую відео...")
    user_id = update.message.from_user.id
    try:
        video   = update.message.video or update.message.video_note
        caption = (update.message.caption if update.message.video else None) or "Що відбувається у цьому відео?"
        if video.file_size > MAX_VIDEO_SIZE:
            await msg.edit_text("❌ Відео завелике. Максимум 20MB.")
            return
        video_bytes = bytes(await (await ctx.bot.get_file(video.file_id)).download_as_bytearray())
        result      = await analyze_video(video_bytes, caption, user_id)
        last_context[user_id] = {"type": "відео", "description": result[:CTX_DESCRIPTION_LEN]}
        await append_and_trim(user_id, "user", f"[Відео] {caption}")
        await append_and_trim(user_id, "assistant", result)
        await _send_or_edit(msg, clean_markdown(result))
    except Exception as e:
        log.error("handle_video: %s", e)
        await msg.edit_text("Помилка аналізу відео. Спробуй ще раз.")

async def handle_sticker(update: Update, ctx: ContextTypes.DEFAULT_TYPE):
    import random
    user_id = update.message.from_user.id
    reaction = random.choice(STICKER_REACTIONS)
    mem  = get_user_memory(user_id)
    name = mem.get("name")
    if name:
        reaction = reaction.replace("сер", name)
    await update.message.reply_text(reaction)

async def handle_document(update: Update, ctx: ContextTypes.DEFAULT_TYPE):
    doc     = update.message.document
    fname   = doc.file_name.lower()
    user_id = update.message.from_user.id
    caption = update.message.caption or "Стисло підсумуй цей документ українською."

    if fname.endswith(".pdf"):
        msg, extractor = await update.message.reply_text("📄 Читаю PDF..."), extract_pdf_text
    elif fname.endswith((".xlsx", ".xls")):
        msg, extractor = await update.message.reply_text("📊 Читаю Excel..."), extract_excel_text
    elif fname.endswith((".docx", ".doc")):
        msg, extractor = await update.message.reply_text("📝 Читаю Word..."), extract_word_text
    elif fname.endswith(".txt"):
        msg, extractor = await update.message.reply_text("📄 Читаю TXT..."), extract_txt_text
    else:
        await update.message.reply_text("📄 Підтримуються: PDF, Excel, Word, TXT")
        return

    raw  = bytes(await (await ctx.bot.get_file(doc.file_id)).download_as_bytearray())
    text = extractor(raw)
    try:
        if not text or text.startswith("Помилка"):
            await msg.edit_text(f"❌ {text}")
            return
        text_preview = text[:MAX_DOC_PREVIEW_CHARS] + ("..." if len(text) > MAX_DOC_PREVIEW_CHARS else "")
        await append_and_trim(user_id, "user", f"{caption}\n\nДокумент:\n{text_preview}")
        reply = await call_ai(get_history(user_id))
        await append_and_trim(user_id, "assistant", reply)
        last_context[user_id] = {"type": "документ", "description": reply[:CTX_DESCRIPTION_LEN]}
        await _send_or_edit(msg, f"📄 {doc.file_name}\n\n{clean_markdown(reply)}")
    except Exception as e:
        log.error("handle_document: %s", e)
        await msg.edit_text("Помилка обробки документу. Спробуй ще раз.")

async def handle_voice(update: Update, ctx: ContextTypes.DEFAULT_TYPE):
    msg = await update.message.reply_text("🎤 Розпізнаю...")
    try:
        file        = await ctx.bot.get_file(update.message.voice.file_id)
        audio_bytes = bytes(await file.download_as_bytearray())
        text        = await transcribe_voice(audio_bytes)
        if not text:
            await msg.edit_text("Не вдалось розпізнати 😔")
            return
        await msg.edit_text(f"🎤 Ти сказав: {text}\n\n⏳ Обробляю...")
        user_id = update.message.from_user.id
        if get_gender(user_id) is None:
            gender = await detect_gender_from_transcript(text)
            if gender:
                update_user_memory(user_id, {"gender": gender})
                if user_id in chat_histories:
                    chat_histories[user_id][0] = get_system_prompt(user_id)
        await _process_message(update, ctx, user_id, text, voice_msg=msg)
    except Exception as e:
        log.error("handle_voice: %s", e)
        await msg.edit_text("Помилка голосового. Спробуй ще раз.")

# ── Startup ───────────────────────────────────────────────────────────────────

async def post_init(app) -> None:
    init_db()
    await restore_reminders(app.bot)
    for uid_str, data in ({} if True else {}).items():  # memory вже в DB
        pass
    # Відновлюємо режими з DB
    with _db() as con:
        rows = con.execute("SELECT user_id, data FROM memory").fetchall()
    for row in rows:
        data = json.loads(row["data"])
        if data.get("mode"):
            user_personalities[row["user_id"]] = data["mode"]
    # Запускаємо планувальник брифу
    asyncio.create_task(_brief_scheduler(app.bot))
    asyncio.create_task(_proactive_scheduler(app.bot))
    log.info("Bot initialized, DB=%s", DB_PATH)

if __name__ == "__main__":
    app = ApplicationBuilder().token(TELEGRAM_TOKEN).post_init(post_init).build()

    # Commands
    app.add_handler(CommandHandler("start",     start))
    app.add_handler(CommandHandler("help",      help_cmd))
    app.add_handler(CommandHandler("mode",      mode_cmd))
    app.add_handler(CommandHandler("memory",    memory_cmd))
    app.add_handler(CommandHandler("forget",    forget_cmd))
    app.add_handler(CommandHandler("reset",     reset))
    app.add_handler(CommandHandler("search",    handle_search))
    app.add_handler(CommandHandler("status",    handle_status))
    app.add_handler(CommandHandler("remind",    handle_remind))
    app.add_handler(CommandHandler("image",     handle_image_cmd))
    app.add_handler(CommandHandler("translate", translate_cmd))
    app.add_handler(CommandHandler("summarize", summarize_cmd))
    app.add_handler(CommandHandler("generate",  generate_cmd))
    app.add_handler(CommandHandler("edit",      edit_cmd))
    app.add_handler(CommandHandler("recipe",    recipe_cmd))
    app.add_handler(CommandHandler("tasks",     handle_tasks_cmd))
    app.add_handler(CommandHandler("brief",     handle_brief_cmd))

    # Callbacks
    app.add_handler(CallbackQueryHandler(handle_mode_callback, pattern="^mode_"))
    app.add_handler(CallbackQueryHandler(handle_task_callback, pattern="^task_"))

    # Media
    app.add_handler(MessageHandler(filters.PHOTO,                        handle_photo))
    app.add_handler(MessageHandler(filters.VOICE,                        handle_voice))
    app.add_handler(MessageHandler(filters.VIDEO | filters.VIDEO_NOTE,   handle_video))
    app.add_handler(MessageHandler(filters.Sticker.ALL,                  handle_sticker))
    app.add_handler(MessageHandler(
        filters.Document.PDF |
        filters.Document.MimeType("application/vnd.openxmlformats-officedocument.spreadsheetml.sheet") |
        filters.Document.MimeType("application/vnd.openxmlformats-officedocument.wordprocessingml.document") |
        filters.Document.MimeType("application/msword") |
        filters.Document.MimeType("application/vnd.ms-excel") |
        filters.Document.MimeType("text/plain"),
        handle_document,
    ))
    app.add_handler(MessageHandler(filters.TEXT & ~filters.COMMAND, handle_message))

    log.info("Bot started")
    app.run_polling(drop_pending_updates=True, allowed_updates=Update.ALL_TYPES)
